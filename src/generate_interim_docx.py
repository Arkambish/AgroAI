"""Generate the FYP Interim Report as a .docx that conforms to the
University of Moratuwa Faculty of IT guidelines (Karunananda 2006).

Produces: outputs/Interim_Report_AgroAI.docx

Layout per guideline:
- A4, Times New Roman 12pt body, 1.5 line spacing
- Margins: left 1.5", top/right/bottom 1"
- Cover page (no page number) + Title page (no page number)
- Pre-pages with Roman numerals (Abstract, ToC, List of Figures/Tables)
- Body in Arabic numerals starting at 1
- Chapter headings 18pt bold, section/subsection headings 12pt bold
- Each chapter has Introduction (1.x.1) and Summary (1.x.N) subsections
- Citations are Karunananda-style: [n] in text, alphabetical list at end
"""

from __future__ import annotations

import os
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(ROOT, "outputs", "plots", "figures")
EDA_DIR = os.path.join(ROOT, "outputs", "plots", "eda")
RESULTS_PLOTS_DIR = os.path.join(ROOT, "outputs", "plots", "results")
TRAINING_DIR = os.path.join(ROOT, "outputs", "plots", "training")
OUT_PATH = os.path.join(ROOT, "outputs", "Interim_Report_AgroAI.docx")

# ---------- Style helpers ----------------------------------------------------

BODY_FONT = "Times New Roman"


def _add_field(paragraph, instr_text: str) -> None:
    """Insert a Word field (e.g. PAGE) into the given paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_char_begin)
    r.append(instr)
    r.append(fld_char_separate)
    r.append(fld_char_end)


def _set_section_page_numbering(section, *, fmt: str, start: int | None = None) -> None:
    """Set page-number format on a section. fmt: 'decimal' | 'lowerRoman' | 'none'."""
    sect_pr = section._sectPr
    # Remove existing pgNumType if any
    for existing in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def _configure_document(doc: Document) -> None:
    # Page size + margins
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # Different first page so cover/title pages can suppress the page number.
    section.different_first_page_header_footer = True

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)


def _add_centered_runs(doc: Document, items: Iterable[tuple[str, int, bool]]) -> None:
    """items: (text, font_pt, bold). Adds a centered paragraph per tuple."""
    for text, size, bold in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(size)
        run.font.bold = bold


def _spacer(doc: Document, blank_lines: int = 1) -> None:
    for _ in range(blank_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def _heading(doc: Document, text: str, *, level: str) -> None:
    """level: 'chapter' (18pt bold), 'section' (12pt bold), 'subsection' (12pt bold)."""
    p = doc.add_paragraph()
    if level == "chapter":
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = BODY_FONT
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = BODY_FONT


def _para(doc: Document, text: str, *, justify: bool = True) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(12)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(12)


def _figure(doc: Document, image_path: str, caption: str, *, width_inches: float = 5.5) -> None:
    if not os.path.exists(image_path):
        # Placeholder so missing assets don't crash the run
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Figure asset missing: {image_path}]")
        run.italic = True
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.name = BODY_FONT
    cap_run.font.size = Pt(11)
    cap_run.italic = True
    cap.paragraph_format.space_after = Pt(12)


def _placeholder(doc: Document, label: str) -> None:
    """A boxed placeholder for a screenshot the user will insert later."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Insert screenshot here — {label} ]")
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    # Light grey divider above and below
    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _section_break(doc: Document) -> None:
    """Add a hard section break so we can change page-number format."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_height = Inches(11.69)
    new_section.page_width = Inches(8.27)
    new_section.left_margin = Inches(1.5)
    new_section.right_margin = Inches(1.0)
    new_section.top_margin = Inches(1.0)
    new_section.bottom_margin = Inches(1.0)


def _add_centered_footer_pagenum(section, fmt: str, start: int | None = None) -> None:
    """Add a centered PAGE field to the section footer and set the format."""
    _set_section_page_numbering(section, fmt=fmt, start=start)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""  # clear
    _add_field(p, "PAGE")


# ---------- Cover and title pages -------------------------------------------

GROUP_MEMBERS = [
    ("214019K", "Arkam B.H.M."),
    ("214192G", "Sharuja B."),
    ("214193K", "Shathurya P."),
]

PROJECT_TITLE_LINE_1 = "AI-Powered Harvest Yield Prediction"
PROJECT_TITLE_LINE_2 = "for Non-Cash Crops (Big Onion) in Sri Lanka"


def _add_member_table(doc: Document) -> None:
    table = doc.add_table(rows=len(GROUP_MEMBERS), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (idx, name) in enumerate(GROUP_MEMBERS):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.text = idx
        c1.text = name
        for c in (c0, c1):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = BODY_FONT
                    run.font.size = Pt(12)
    # Remove all borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "nil")
                tc_borders.append(b)
            tc_pr.append(tc_borders)


def add_cover_page(doc: Document) -> None:
    _spacer(doc, 3)
    _add_centered_runs(doc, [("Interim Report", 16, True)])
    _spacer(doc, 1)
    _add_centered_runs(doc, [("Level 4", 14, True)])
    _spacer(doc, 4)
    _add_centered_runs(
        doc,
        [
            (PROJECT_TITLE_LINE_1, 14, True),
            (PROJECT_TITLE_LINE_2, 14, True),
        ],
    )
    _spacer(doc, 5)
    _add_centered_runs(doc, [("Group Name: Agro AI", 12, False)])
    _spacer(doc, 1)
    _add_member_table(doc)
    _spacer(doc, 6)
    _add_centered_runs(
        doc,
        [
            ("Faculty of Information Technology", 12, False),
            ("University of Moratuwa", 12, False),
            ("2026", 12, False),
        ],
    )
    _page_break(doc)


def add_title_page(doc: Document) -> None:
    _spacer(doc, 3)
    _add_centered_runs(doc, [("Interim Report", 16, True)])
    _spacer(doc, 1)
    _add_centered_runs(doc, [("Level 4", 14, True)])
    _spacer(doc, 4)
    _add_centered_runs(
        doc,
        [
            (PROJECT_TITLE_LINE_1, 14, True),
            (PROJECT_TITLE_LINE_2, 14, True),
        ],
    )
    _spacer(doc, 4)
    _add_centered_runs(doc, [("Group Name: Agro AI", 12, False)])
    _spacer(doc, 1)
    _add_member_table(doc)
    _spacer(doc, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("Supervised by: ")
    r1.font.bold = True
    r1.font.name = BODY_FONT
    r1.font.size = Pt(12)
    r2 = p.add_run("Dr. Firdhous M.F.M.")
    r2.font.name = BODY_FONT
    r2.font.size = Pt(12)
    _spacer(doc, 4)
    _add_centered_runs(
        doc,
        [
            ("Faculty of Information Technology", 12, False),
            ("University of Moratuwa", 12, False),
            ("2026", 12, False),
        ],
    )


# ---------- Chapter content (text only) -------------------------------------

# To keep this file readable, the long-form chapter prose lives in CONTENT below.
# Each chapter is a list of (kind, payload) tuples:
#   ("h2", "1.1 Introduction")
#   ("h3", "1.1.1 Sub-section")
#   ("p", "Long paragraph...")
#   ("bullet", "list item")
#   ("fig", ("path", "Figure 1.1: Caption"))
#   ("ph", "screenshot label")           # screenshot placeholder
#   ("table", (header_list, [row1, row2, ...], "Table 2.1: Caption"))


def _add_table(doc: Document, header: list[str], rows: list[list[str]], caption: str) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
    # Body
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.name = BODY_FONT
    cr.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(12)


def _emit(doc: Document, items: list[tuple]) -> None:
    for kind, payload in items:
        if kind == "h1":
            _heading(doc, payload, level="chapter")
        elif kind == "h2":
            _heading(doc, payload, level="section")
        elif kind == "h3":
            _heading(doc, payload, level="subsection")
        elif kind == "p":
            _para(doc, payload)
        elif kind == "bullet":
            _bullet(doc, payload)
        elif kind == "fig":
            path, caption = payload
            _figure(doc, path, caption)
        elif kind == "ph":
            _placeholder(doc, payload)
        elif kind == "table":
            header, rows, caption = payload
            _add_table(doc, header, rows, caption)
        elif kind == "pagebreak":
            _page_break(doc)
        else:
            raise ValueError(f"Unknown emit kind: {kind}")


# ---------- Pre-pages -------------------------------------------------------

ABSTRACT_TEXT = (
    "Sri Lanka consumes approximately 220,000 metric tons of big onion "
    "(Allium cepa) annually, but domestic production covers only a fraction of "
    "this demand and the balance is imported at significant foreign-exchange "
    "cost. Unlike paddy, no systematic crop-cutting survey methodology exists "
    "for big onion, so yield estimates rely on subjective field assessments "
    "by agricultural officers and arrive only after harvest. This project "
    "addresses that information gap by designing, implementing, and evaluating "
    "an artificial-intelligence-powered yield prediction system for big onion "
    "in the four major producing districts of Sri Lanka: Matale, Anuradhapura, "
    "Polonnaruwa, and Jaffna, across both the Yala and Maha cultivation seasons. "
    "The proposed system integrates four data streams — historical Department "
    "of Census and Statistics yield records, daily NASA POWER weather data, "
    "satellite-derived vegetation indices from MODIS and Sentinel-2 obtained via "
    "Google Earth Engine, and SoilGrids soil characteristics — and trains a "
    "comparative suite of three machine learning models (Random Forest, "
    "XGBoost, Support Vector Regression) and four deep learning models "
    "(LSTM, Bidirectional LSTM, one-dimensional Convolutional Neural Network, "
    "and a novel Hybrid CNN-LSTM architecture with explicit season-indicator "
    "injection). Models are evaluated using Leave-One-Year-Out cross-validation "
    "with statistical significance testing (Wilcoxon and paired t-tests). "
    "Predictions, model comparison metrics, and SHAP-based feature attribution "
    "are exposed through a Flask REST API and consumed by a Next.js decision-support "
    "dashboard with an interactive choropleth, smart prediction form, and admin "
    "panel. The dashboard targets the Department of Census and Statistics, the "
    "Department of Agriculture, district agricultural officers, and researchers. "
    "The interim implementation runs end-to-end on a synthetic dataset calibrated "
    "to published yield distributions, with the architecture validated and ready "
    "for ingestion of real survey and satellite data."
)


def add_pre_pages(doc: Document) -> None:
    # Abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = BODY_FONT
    p.paragraph_format.space_after = Pt(18)
    _para(doc, ABSTRACT_TEXT)
    _page_break(doc)

    # Acknowledgements
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Acknowledgements")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = BODY_FONT
    p.paragraph_format.space_after = Pt(18)
    _para(
        doc,
        "We extend our sincere gratitude to our supervisor, Dr. Firdhous M.F.M., "
        "Faculty of Information Technology, University of Moratuwa, for his "
        "guidance and continued support throughout the project. We are grateful "
        "to the Department of Census and Statistics of Sri Lanka for making "
        "the big onion survey reports publicly available, to NASA POWER and "
        "the Climate Hazards Group for free access to weather and rainfall "
        "datasets, to the European Space Agency Copernicus programme and the "
        "MODIS Science Team for satellite imagery, and to the ISRIC SoilGrids "
        "team for global soil property maps. We thank the Faculty of "
        "Information Technology for providing the academic environment that "
        "made this work possible.",
    )
    _page_break(doc)

    # Table of Contents (manual — Word can regenerate this on open if asked)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Contents")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = BODY_FONT
    p.paragraph_format.space_after = Pt(18)
    toc_entries = [
        ("Abstract", "i"),
        ("Acknowledgements", "ii"),
        ("Contents", "iii"),
        ("List of Figures", "v"),
        ("List of Tables", "vi"),
        ("Chapter 1 — Introduction", "1"),
        ("    1.1 Introduction", "1"),
        ("    1.2 Background and Motivation", "3"),
        ("    1.3 Problem in Brief", "5"),
        ("    1.4 Aim and Objectives", "6"),
        ("    1.5 Proposed Solution Overview", "7"),
        ("    1.6 Structure of the Report", "8"),
        ("    1.7 Summary", "9"),
        ("Chapter 2 — Crop Yield Prediction: State of the Art", "10"),
        ("    2.1 Introduction", "10"),
        ("    2.2 Machine Learning Approaches", "11"),
        ("    2.3 Deep Learning Approaches", "13"),
        ("    2.4 Hybrid Architectures", "15"),
        ("    2.5 Crop Yield Prediction in Sri Lanka", "17"),
        ("    2.6 Onion-Specific Research", "18"),
        ("    2.7 Research Gaps Identified", "19"),
        ("    2.8 Comparison of Existing Approaches", "20"),
        ("    2.9 Summary", "21"),
        ("Chapter 3 — Technologies Adopted", "22"),
        ("    3.1 Introduction", "22"),
        ("    3.2 Data Sources and Collection Technologies", "22"),
        ("    3.3 Machine Learning Technologies", "24"),
        ("    3.4 Deep Learning Technologies", "26"),
        ("    3.5 Remote Sensing and Geospatial Technologies", "27"),
        ("    3.6 Software and Tools", "28"),
        ("    3.7 Evaluation Methodology", "29"),
        ("    3.8 Summary", "30"),
        ("Chapter 4 — System Approach", "31"),
        ("    4.1 Introduction", "31"),
        ("    4.2 Users", "31"),
        ("    4.3 Inputs and Outputs", "32"),
        ("    4.4 System Pipeline", "33"),
        ("    4.5 High-Level System Architecture", "35"),
        ("    4.6 End-to-End Walkthrough of a Single Prediction", "36"),
        ("    4.7 Non-Functional Requirements", "37"),
        ("    4.8 Summary", "38"),
        ("Chapter 5 — Analysis and Design", "39"),
        ("    5.1 Introduction", "39"),
        ("    5.2 Data Flow Design", "39"),
        ("    5.3 Feature Engineering Design", "41"),
        ("    5.4 ML Model Design", "43"),
        ("    5.5 DL Model Design and Hybrid CNN-LSTM Novelty", "45"),
        ("    5.6 Data Pipeline Design", "47"),
        ("    5.7 Dashboard Design", "48"),
        ("    5.8 Database Schema", "49"),
        ("    5.9 Summary", "50"),
        ("Chapter 6 — Implementation", "51"),
        ("    6.1 Introduction", "51"),
        ("    6.2 Repository Layout and Tooling", "51"),
        ("    6.3 Data Pipeline Implementation", "52"),
        ("    6.4 ML Model Implementation", "53"),
        ("    6.5 DL Model Implementation", "54"),
        ("    6.6 Ablation Study and SHAP", "55"),
        ("    6.7 Flask REST API", "56"),
        ("    6.8 Dashboard Implementation", "57"),
        ("    6.9 Preliminary Results on Synthetic Data", "58"),
        ("    6.10 Summary", "59"),
        ("Chapter 7 — Discussion and Further Work", "60"),
        ("    7.1 Introduction", "60"),
        ("    7.2 How This Work Differs from Others", "60"),
        ("    7.3 Challenges Encountered", "62"),
        ("    7.4 Further Work Plan", "63"),
        ("    7.5 Methodological Choices and Their Justifications", "64"),
        ("    7.6 Limitations and Risks", "66"),
        ("    7.7 Summary", "67"),
        ("References", "68"),
        ("Appendix A — Individuals' Contribution to the Project", "70"),
        ("Appendix B — Code and Implementation Excerpts", "73"),
    ]
    for label, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
        run = p.add_run(label)
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
        if not label.startswith("    "):
            run.font.bold = True
        p.add_run("\t").font.name = BODY_FONT
        page_run = p.add_run(page)
        page_run.font.name = BODY_FONT
        page_run.font.size = Pt(12)
        if not label.startswith("    "):
            page_run.font.bold = True
        p.paragraph_format.space_after = Pt(2)
    _page_break(doc)

    # List of Figures
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("List of Figures")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = BODY_FONT
    p.paragraph_format.space_after = Pt(18)
    figures = [
        ("Figure 4.1 — High-level system architecture", "35"),
        ("Figure 4.2 — System pipeline stages", "34"),
        ("Figure 5.1 — Data flow diagram", "40"),
        ("Figure 5.2 — Hybrid CNN-LSTM architecture (novelty)", "46"),
        ("Figure 5.3 — Database schema", "49"),
        ("Figure 6.1 — Yield distribution by district and season", "52"),
        ("Figure 6.2 — Correlation heatmap of features with yield", "53"),
        ("Figure 6.3 — Model comparison: RMSE and R²", "58"),
        ("Figure 6.4 — Ablation study: feature-source contribution", "59"),
        ("Figure 6.5 — SHAP top-15 feature importance", "59"),
    ]
    for label, page in figures:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
        run = p.add_run(label)
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
        p.add_run("\t").font.name = BODY_FONT
        page_run = p.add_run(page)
        page_run.font.name = BODY_FONT
        page_run.font.size = Pt(12)
    _page_break(doc)

    # List of Tables
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("List of Tables")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = BODY_FONT
    p.paragraph_format.space_after = Pt(18)
    tables = [
        ("Table 2.1 — Comparison of existing approaches", "20"),
        ("Table 3.1 — Data sources and details", "23"),
        ("Table 3.2 — Software and tools", "29"),
        ("Table 5.1 — Feature engineering summary", "42"),
        ("Table 6.1 — Preliminary model comparison on synthetic data", "58"),
        ("Table 6.2 — Ablation experiment results", "59"),
        ("Table A.1 — Distribution of individual contributions", "72"),
    ]
    for label, page in tables:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
        run = p.add_run(label)
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
        p.add_run("\t").font.name = BODY_FONT
        page_run = p.add_run(page)
        page_run.font.name = BODY_FONT
        page_run.font.size = Pt(12)


# ---------- Body chapters ---------------------------------------------------

# (To keep this file at a reasonable length, the chapter content is
# defined as data in the get_*_chapter functions below.)


def get_chapter_1() -> list[tuple]:
    return [
        ("h1", "Chapter 1 — Introduction"),
        ("h2", "1.1 Introduction"),
        ("p",
         "Agriculture remains the backbone of Sri Lanka's rural economy. "
         "Among non-cash crops, big onion (Allium cepa) holds a "
         "strategically important position: Sri Lanka consumes between "
         "200,000 and 220,000 metric tons annually, yet domestic production "
         "covers only a fraction of this demand, with the balance met "
         "through costly imports that drain foreign exchange reserves [6, 8]."),
        ("p",
         "Accurate and timely yield prediction is essential for "
         "agricultural planning, import regulation, and food security "
         "policy. However, current estimation methods for vegetable crops "
         "rely on manual surveys and post-harvest assessments [4]. Unlike "
         "paddy, which is surveyed through crop-cutting surveys over four "
         "thousand tracts per season, no equivalent methodology exists for "
         "vegetables — leaving forecasting an exercise in educated "
         "estimation rather than data-driven analysis."),
        ("p",
         "Recent advances in machine learning (ML) and deep learning (DL), "
         "combined with the growing availability of open-access satellite "
         "and climate datasets, present an opportunity to develop cost-"
         "effective data-driven yield prediction systems [1, 2]. ML and DL "
         "have been applied successfully to cereal crops globally [3, 7], "
         "but applications to vegetables remain under-explored, and no "
         "published study has applied ML or DL to big onion yield "
         "prediction in Sri Lanka. This project addresses that gap by "
         "proposing an AI-powered yield prediction system that integrates "
         "multi-source data for timely, scalable, and explainable forecasts."),

        ("h2", "1.2 Background and Motivation"),
        ("p",
         "Big onion cultivation in Sri Lanka is concentrated in Matale, "
         "Anuradhapura, Polonnaruwa, Puttalam, and Jaffna, with the Matale "
         "District (particularly Dambulla) accounting for over fifty percent "
         "of national production [6]. Cultivation follows the country's two "
         "monsoon-driven seasons: Yala (April–August), the primary production "
         "period in the dry zones, and Maha (October–March), the off-season. "
         "Yield variation between seasons is substantial — peak-season yields "
         "in Matale average approximately 8,800 kg/acre compared to off-"
         "season yields of around 3,400 kg/acre [4]. This twofold seasonal "
         "differential makes the bimodal monsoon system a defining feature "
         "of any predictive model."),
        ("p",
         "The motivation for this research stems from four factors. First, "
         "the strategic economic importance of big onion and the government's "
         "self-sufficiency objective make accurate forecasting a national "
         "priority. Second, open-access satellite data (MODIS, Sentinel-2, "
         "CHIRPS) and climate datasets (NASA POWER) make the data "
         "infrastructure available at no cost [3, 12]. Third, the success "
         "of AI techniques in agricultural forecasting globally [7, 8] "
         "indicates the methodological foundations are mature. Fourth, no "
         "existing AI-based system for big onion yield prediction in Sri "
         "Lanka exists, highlighting a clear research gap."),

        ("h2", "1.3 Problem in Brief"),
        ("p",
         "Yield information for big onion is currently available only during "
         "or after harvest, preventing farmers, policymakers, and import "
         "regulators from making proactive decisions. Existing methods rely "
         "on subjective assessments and do not incorporate environmental "
         "data such as satellite vegetation indices, climate records, and "
         "soil characteristics — all of which critically influence yield. "
         "Agricultural authorities therefore face challenges in import "
         "planning, while farmers lack early yield guidance that could "
         "inform resource allocation and marketing. The system developed "
         "here provides accurate early-season forecasts across the major "
         "producing districts, supporting informed decisions for all "
         "stakeholders."),

        ("h2", "1.4 Aim and Objectives"),
        ("p",
         "The aim of this project is to design, develop, and evaluate an AI-"
         "powered system for predicting big onion harvest yield in the four "
         "major producing districts of Sri Lanka — Matale, Anuradhapura, "
         "Polonnaruwa, and Jaffna — using machine learning, deep learning, "
         "and remote sensing technologies."),
        ("p", "The specific objectives are as follows:"),
        ("bullet",
         "To study big onion cultivation patterns, seasonal dynamics, and "
         "yield-influencing factors in Sri Lanka."),
        ("bullet",
         "To collect, integrate, and preprocess multi-source datasets "
         "including DCS yield records, NASA POWER weather, MODIS and "
         "Sentinel-2 vegetation indices via Google Earth Engine, and "
         "SoilGrids soil characteristics."),
        ("bullet",
         "To design and implement machine learning models (Random Forest, "
         "XGBoost, SVR) and deep learning models (LSTM, BiLSTM, 1D CNN, "
         "and a novel Hybrid CNN-LSTM)."),
        ("bullet",
         "To evaluate and compare model accuracy using RMSE, MAE, R², and "
         "MAPE under Leave-One-Year-Out cross-validation."),
        ("bullet",
         "To conduct a quantitative ablation study decomposing the "
         "contribution of each data source."),
        ("bullet",
         "To apply SHAP-based feature attribution to the best model for "
         "interpretable per-feature explanations."),
        ("bullet",
         "To develop a decision-support dashboard with district-level "
         "spatial yield maps and seasonal forecasts."),
        ("bullet",
         "To prepare reproducible documentation suitable for hand-off to "
         "agricultural agencies."),

        ("h2", "1.5 Proposed Solution Overview"),
        ("p",
         "The proposed solution comprises four components: (1) an automated "
         "data pipeline that integrates satellite imagery, weather, yield, "
         "and soil data; (2) a feature engineering module that computes "
         "vegetation indices, growing degree days, and drought indices; "
         "(3) a predictive modelling engine that trains and compares ML "
         "and DL models under LOYO-CV; and (4) a visualisation dashboard "
         "presenting district-level predictions and SHAP attributions "
         "through a web interface. Trained models are served through a "
         "Flask REST API."),
        ("p",
         "End users include the Department of Census and Statistics, the "
         "Department of Agriculture, district agricultural officers, and "
         "researchers. Inputs are weather variables, satellite vegetation "
         "indices, historical yield, and soil characteristics. Outputs are "
         "predicted yield values in metric tons per hectare with confidence "
         "intervals, district-level spatial maps, seasonal comparisons, and "
         "SHAP feature attributions."),

        ("h2", "1.6 Structure of the Report"),
        ("p",
         "Chapter 2 reviews existing research on crop yield prediction and "
         "identifies research gaps. Chapter 3 describes the technologies "
         "adopted. Chapter 4 presents the system approach. Chapter 5 "
         "provides the detailed analysis and design including the novel "
         "Hybrid CNN-LSTM architecture. Chapter 6 reports on implementation "
         "progress and preliminary results on a calibrated synthetic "
         "dataset. Chapter 7 discusses how the work differs from existing "
         "approaches, challenges encountered, and the plan for the remaining "
         "project period. References (alphabetised) and Appendices "
         "(individual contributions, code excerpts) follow."),

        ("h2", "1.7 Summary"),
        ("p",
         "This chapter has introduced the problem of inaccurate and delayed "
         "big onion yield estimation in Sri Lanka, the motivation for an AI-"
         "based, multi-source approach, and the project's aim and "
         "objectives. The next chapter reviews existing literature on crop "
         "yield prediction using ML and DL techniques and identifies the "
         "research gaps that this project addresses."),
        ("pagebreak", None),
    ]


def get_chapter_2() -> list[tuple]:
    return [
        ("h1", "Chapter 2 — Crop Yield Prediction: State of the Art"),
        ("h2", "2.1 Introduction"),
        ("p",
         "The previous chapter introduced the problem of big onion yield "
         "estimation in Sri Lanka and outlined the proposed AI-based "
         "solution. This chapter reviews existing research on crop yield "
         "prediction using machine learning and deep learning techniques. "
         "The review is organised in four parts: machine learning approaches "
         "for tabular yield prediction, deep learning approaches that exploit "
         "temporal and spatial structure, hybrid architectures combining "
         "spatial and temporal modalities, and Sri Lankan and onion-specific "
         "studies. The chapter concludes by identifying the research gaps "
         "that this project addresses and presenting a comparison of "
         "existing approaches."),

        ("h2", "2.2 Machine Learning Approaches for Crop Yield Prediction"),
        ("p",
         "Jabed et al. [1] reviewed 115 papers on ML-based yield "
         "prediction and reported that Random Forest, SVM, and Artificial "
         "Neural Networks are the most commonly used algorithms, with "
         "temperature, soil type, and vegetation indices as the most "
         "utilised features. The review also found that models with larger "
         "feature sets do not consistently outperform models with "
         "carefully selected features — feature engineering driven by "
         "agronomic insight is often more impactful than model complexity."),
        ("p",
         "Chikwendu et al. [3] compared Random Forest, 1D CNN, Gradient "
         "Boosting, and Decision Trees for potato, maize, and cotton, "
         "finding Random Forest achieved R² = 0.875 for potatoes while "
         "XGBoost showed the lowest error for cotton. The cross-crop "
         "variability reinforces the importance of empirical comparison "
         "rather than blind adoption of a single algorithm — a principle "
         "this project follows by training seven distinct models."),
        ("p",
         "Random Forest is a strong baseline for tabular agricultural data: "
         "it is robust to outliers, scale-invariant, captures non-linear "
         "interactions through tree splits, and provides feature importance "
         "[1, 17]. XGBoost extends gradient boosting with second-order "
         "gradient information and built-in regularisation [18]. Support "
         "Vector Regression provides a theoretically grounded kernel-based "
         "alternative. A common limitation in agriculture is small training-"
         "set size: vegetable datasets in Sri Lanka contain fewer than two "
         "hundred records, conditions under which classical ML often "
         "outperforms deep models."),

        ("h2", "2.3 Deep Learning Approaches"),
        ("p",
         "Deep learning methods, particularly LSTMs and CNNs, have shown "
         "promise for capturing temporal and spatial patterns in "
         "agricultural data. Wang et al. [7] compared an LSTM-with-"
         "attention model against XGBoost and Random Forest for maize "
         "yield in the US Corn Belt, reporting that the LSTM explained "
         "73 % of the spatiotemporal variance — outperforming the "
         "ensemble baselines. The attention mechanism enabled the model "
         "to weight monthly weather observations differently by growth "
         "stage."),
        ("p",
         "LSTMs are well suited to weather-based yield prediction because "
         "the cultivation season unfolds as a sequence of monthly "
         "observations, and the gating architecture mitigates the "
         "vanishing-gradient problem of vanilla recurrent networks [16]. "
         "Bidirectional LSTMs augment this by processing the sequence "
         "both forwards and backwards, exposing each timestep to past and "
         "future context within the season."),
        ("p",
         "Convolutional Neural Networks have been applied to satellite "
         "imagery in two ways: 2D CNNs extract spatial features from "
         "raster vegetation maps, while 1D CNNs operate on aggregated "
         "feature vectors and learn local interactions through small "
         "filters [20]. The 1D approach is used in this project, where "
         "satellite indices are summarised into seasonal statistics."),

        ("h2", "2.4 Hybrid Architectures Combining Spatial and Temporal Modalities"),
        ("p",
         "Hybrid CNN-LSTM architectures combine convolutional spatial "
         "feature extraction with recurrent temporal modelling. Rajpoot "
         "and Chandrakar [8] presented a hybrid CNN-LSTM framework "
         "achieving R² = 0.92 on a general crop yield benchmark, "
         "outperforming standalone LSTM, CNN, RF, and SVR on the same "
         "dataset. NDVI, rainfall, and temperature were identified as "
         "the most influential features."),
        ("p",
         "Hybrid architectures follow one of two patterns: sequential "
         "(CNN → LSTM, treating the CNN output as a sequence) or "
         "parallel-branch (CNN and LSTM operate on different inputs and "
         "their outputs are concatenated before dense layers). The "
         "parallel-branch pattern is more flexible because each branch "
         "specialises on the modality best suited to it. The novel "
         "architecture in this project follows this pattern and injects "
         "an explicit season-indicator scalar into the merged "
         "representation, detailed in Chapter 5."),
        ("p",
         "Transformer-based models have also been applied to crop yield, "
         "but their parameter counts exceed LSTMs by an order of "
         "magnitude, making them ill-suited to the data-scarce vegetable-"
         "yield regime. CNN-LSTM is chosen as the novelty platform "
         "because it is expressive enough to fuse multi-modal inputs "
         "while remaining trainable on the ~160 observations available."),

        ("h2", "2.5 Crop Yield Prediction in Sri Lanka"),
        ("p",
         "Research on AI-based yield prediction in Sri Lanka has focused "
         "almost exclusively on paddy. Amarasinghe et al. [5] used weather-"
         "based feature-engineered ML models for rice yield in two "
         "districts (1982–2019), demonstrating that engineered features "
         "such as cumulative rainfall, drought-period counts, and growing-"
         "degree-day aggregates improve accuracy substantially over raw "
         "weather inputs. They reported that engineered features mattered "
         "more than model choice."),
        ("p",
         "Wickramasinghe et al. [10] modelled rice yield and climate using "
         "statistical and ML techniques, finding rainfall during the "
         "vegetative phase was dominant in the dry zone while solar "
         "radiation mattered more in the wet zone. The geographic "
         "specificity reinforces the need to include district-level "
         "features in any multi-district model. Herath et al. [12] "
         "extracted paddy phenological parameters from 16 years of MODIS "
         "NDVI data, establishing that MODIS imagery is sufficiently "
         "reliable over Sri Lankan terrain — the same source is used "
         "here for big onion vegetation indices."),
        ("p",
         "Despite this paddy-focused body of work, no published study has "
         "applied ML or DL to vegetable yield prediction in Sri Lanka. "
         "This gap is partly explained by the absence of structured DCS "
         "survey data for vegetables. The present project takes the "
         "available DCS data as a starting point and supplements it with "
         "multi-source remote-sensing and weather inputs."),

        ("h2", "2.6 Onion-Specific Research"),
        ("p",
         "Published research on onion yield prediction is extremely "
         "limited globally. Iqbal et al. [9] developed OnionBangla using "
         "Gradient Boosting, XGBoost, and Random Forest with Bangladeshi "
         "climate data, achieving 94 % accuracy on a binary classification "
         "of yield tiers. However, the study used only weather-based "
         "tabular features without satellite vegetation indices and did "
         "not consider deep learning. Kim and Soon [11] predicted onion "
         "bulb weight in Korea using Neural Prophet — a finer-grained "
         "target than aggregate yield, focused on individual-bulb biology "
         "rather than district-level production."),
        ("p",
         "In the Sri Lankan context, only onion price prediction has been "
         "attempted, with autoregressive models, LSTMs, and Random Forests "
         "applied to Jaffna red onion prices [14]. Price prediction does "
         "not require the multi-source environmental modelling that yield "
         "prediction does. To our knowledge, no published study has "
         "applied AI techniques to big onion yield prediction in Sri Lanka."),

        ("h2", "2.7 Research Gaps Identified"),
        ("p",
         "Five research gaps emerge from the literature review:"),
        ("bullet",
         "No AI-based yield prediction system exists for any vegetable "
         "crop in Sri Lanka — all published Sri Lankan work focuses on "
         "paddy."),
        ("bullet",
         "No systematic yield estimation methodology exists for Sri "
         "Lankan vegetables. DCS relies on subjective assessments for big "
         "onion, unlike the crop-cutting methodology used for paddy."),
        ("bullet",
         "Existing Sri Lankan ML/DL studies use single data sources "
         "(typically weather only) [5, 10]. There has been no multi-"
         "source fusion combining satellite, weather, soil, and historical "
         "yield. The present project fuses four sources and quantifies "
         "each contribution through ablation."),
        ("bullet",
         "No district-level spatial yield mapping exists for Sri Lankan "
         "vegetables. Current studies produce only national aggregates."),
        ("bullet",
         "No rigorous ML-versus-DL comparison has been published for "
         "vegetable crops in tropical monsoon climates. Cereal-crop "
         "comparisons exist [3, 7] but the small-data, bimodal-monsoon "
         "regime is qualitatively different and merits empirical "
         "assessment using LOYO-CV and statistical significance testing."),

        ("h2", "2.8 Comparison of Existing Approaches"),
        ("p",
         "Table 2.1 summarises the comparison of existing approaches and "
         "highlights how this project differs from prior work. The summary "
         "table makes explicit that the present project is the first to "
         "combine four data sources, train both ML and DL models, target a "
         "vegetable crop in Sri Lanka, and propose a season-aware hybrid "
         "architecture."),
        ("table", (
            ["Study", "Crop", "Country", "Methods", "Data Sources", "Limitation"],
            [
                ["Amarasinghe et al. [5]", "Rice", "Sri Lanka", "RF, XGBoost", "Weather only", "No satellite, no DL, rice only"],
                ["OnionBangla [9]", "Onion", "Bangladesh", "RF, XGBoost, GB", "Weather only", "No satellite, no DL"],
                ["Wang et al. [7]", "Maize", "USA", "LSTM, XGBoost, RF", "Weather + soil", "No satellite indices"],
                ["Rajpoot & Chandrakar [8]", "General", "Global", "CNN-LSTM hybrid", "Satellite + weather", "Cereal crops, temperate climate"],
                ["Wickramasinghe [10]", "Rice", "Sri Lanka", "Statistical + ML", "Climate", "No satellite, rice only"],
                ["Kim & Soon [11]", "Onion", "Korea", "Neural Prophet", "Time-series", "Bulb weight, not aggregate yield"],
                ["This project", "Big Onion", "Sri Lanka", "RF, XGB, SVR, LSTM, BiLSTM, CNN, Hybrid", "Satellite + weather + soil + historical", "First for vegetable in Sri Lanka"],
            ],
            "Table 2.1: Comparison of existing approaches",
        )),

        ("h2", "2.9 Summary"),
        ("p",
         "This chapter has reviewed the existing crop yield prediction "
         "literature across global ML approaches, deep learning models, "
         "hybrid spatial-temporal architectures, Sri Lankan studies, and "
         "the limited body of onion-specific research. Five research gaps "
         "have been identified, and the comparative analysis in Table 2.1 "
         "shows that no prior work combines multi-source fusion, ML-versus-"
         "DL comparison, and a vegetable-crop focus in the Sri Lankan "
         "context. The next chapter describes the technologies adopted to "
         "address these gaps."),
        ("pagebreak", None),
    ]


def get_chapter_3() -> list[tuple]:
    return [
        ("h1", "Chapter 3 — Technologies Adopted"),
        ("h2", "3.1 Introduction"),
        ("p",
         "The previous chapter identified five research gaps in crop yield "
         "prediction for vegetable crops in tropical climates. This chapter "
         "describes the technologies adopted to address those gaps, "
         "covering the eight data sources from which the system ingests "
         "data, the machine learning and deep learning frameworks used to "
         "train predictive models, the remote sensing and geospatial tools "
         "used to extract satellite imagery, and the visualisation "
         "technologies used to communicate predictions to end users."),

        ("h2", "3.2 Data Sources and Collection Technologies"),
        ("p",
         "The system integrates data from eight sources spanning satellite "
         "imagery, climate observations, official government surveys, and "
         "global soil databases. Table 3.1 summarises each source, its "
         "spatial and temporal resolution, and the access mechanism by "
         "which the project pipeline retrieves it. Several of these sources "
         "are accessed through Google Earth Engine, which provides "
         "planetary-scale on-demand processing of petabyte-scale satellite "
         "archives at no cost for academic use [12]."),
        ("table", (
            ["Data Type", "Source", "Resolution / Detail", "Access Method"],
            [
                ["Historical Yield", "DCS Big Onion Surveys", "District-level, by season", "PDF reports from statistics.gov.lk"],
                ["National Production", "FAOSTAT", "Country-level, 1961–present", "CSV download from fao.org/faostat"],
                ["Weather Data", "NASA POWER API", "0.5° grid, daily", "Free REST API, no key required"],
                ["Rainfall", "CHIRPS via GEE", "5 km, daily", "Google Earth Engine"],
                ["Vegetation Indices", "MODIS MOD13Q1 via GEE", "250 m, 16-day composites", "Google Earth Engine"],
                ["High-Res Imagery", "Sentinel-2 via GEE", "10 m multispectral", "Google Earth Engine"],
                ["Land Surface Temp", "MODIS LST via GEE", "1 km, 8-day", "Google Earth Engine"],
                ["Soil Data", "SoilGrids", "250 m, static", "REST API at rest.isric.org"],
            ],
            "Table 3.1: Data sources and details",
        )),
        ("p",
         "Multi-source integration reflects the principle [1, 7, 8] that "
         "yield is determined by the joint effect of weather, vegetation "
         "development, and soil context. No single source captures all "
         "three: weather APIs provide climate but no vegetation status; "
         "satellite indices report vegetation but no water balance; soil "
         "databases describe long-term properties but no within-season "
         "variation. Multi-source fusion addresses this complementarity."),

        ("h2", "3.3 Machine Learning Technologies"),
        ("p",
         "Three ML algorithms are used as baselines, all implemented with "
         "scikit-learn [21] and XGBoost. The choice reflects complementary "
         "inductive biases: tree ensembles for non-linear tabular data, "
         "gradient boosting for high-accuracy structured prediction, and "
         "kernel methods for smooth non-linear relationships."),
        ("h3", "3.3.1 Random Forest"),
        ("p",
         "Random Forest [17] is a bagging ensemble of decision trees, "
         "each trained on a bootstrap sample with a random subset of "
         "features at each split. This double randomisation decorrelates "
         "trees and reduces variance. RF handles noisy data robustly, "
         "requires no scaling, and provides feature importance via mean "
         "decrease in impurity. The hyperparameter grid covers 100–500 "
         "trees, maximum depth 5 to unlimited, and varying minimum "
         "samples per leaf, searched using TimeSeriesSplit."),
        ("h3", "3.3.2 XGBoost"),
        ("p",
         "XGBoost [18] is a gradient boosting algorithm combining many "
         "weak tree learners through a stage-wise additive model. It "
         "differs from generic gradient boosting in using second-order "
         "gradient information, built-in L1/L2 regularisation, and "
         "automatic missing-value handling. The project tunes learning "
         "rate, max depth, number of estimators, subsample fraction, and "
         "regularisation strengths."),
        ("h3", "3.3.3 Support Vector Regression"),
        ("p",
         "SVR extends the support vector machine framework to continuous "
         "prediction, fitting a function such that most training points "
         "lie within an epsilon-wide tube. Three kernels (RBF, linear, "
         "polynomial) are evaluated. SVR requires feature standardisation, "
         "applied inside each fold to prevent leakage."),

        ("h2", "3.4 Deep Learning Technologies"),
        ("p",
         "Four DL architectures are implemented using TensorFlow 2 / "
         "Keras 3 [22] functional API, exploring complementary inductive "
         "biases for sequence and tabular data."),
        ("h3", "3.4.1 LSTM and Bidirectional LSTM"),
        ("p",
         "LSTMs [16] capture temporal dependencies in weather time-series "
         "over the four-to-five-month cultivation period [7]. The forget, "
         "input, and output gates enable selective retention of relevant "
         "information across timesteps. Bidirectional LSTMs run two LSTMs "
         "in parallel (forward and backward) and concatenate their hidden "
         "states, exposing each timestep to past and future context."),
        ("h3", "3.4.2 One-Dimensional CNN"),
        ("p",
         "1D CNNs [8] extract patterns from satellite vegetation features "
         "aggregated into fixed-length vectors. Small filters slide along "
         "the feature axis to detect local interactions. Batch "
         "normalisation, max pooling, and global average pooling stabilise "
         "training and produce a permutation-invariant representation."),
        ("h3", "3.4.3 Hybrid CNN-LSTM with Season Indicator"),
        ("p",
         "The hybrid CNN-LSTM is the central novel contribution. Two "
         "parallel branches process satellite features (1D CNN) and "
         "weather time-series (stacked LSTM). Their outputs are "
         "concatenated with a scalar season indicator (Yala = 1, Maha = 0) "
         "and fed through dense layers. Injecting the indicator after "
         "feature extraction lets the CNN/LSTM branches learn season-"
         "agnostic extractors while the final dense layers learn season-"
         "specific yield baselines. The architecture is detailed in "
         "Chapter 5."),

        ("h2", "3.5 Remote Sensing and Geospatial Technologies"),
        ("p",
         "Google Earth Engine [23] provides on-demand processing of "
         "satellite data. The project uses the GEE JavaScript API for "
         "exploratory analysis and the Python API (earthengine-api) for "
         "scriptable extraction. Three vegetation indices are computed:"),
        ("bullet",
         "NDVI = (NIR − Red) / (NIR + Red), measuring vegetation "
         "greenness on a −1 to +1 scale. Healthy crops typically show "
         "NDVI 0.6–0.9."),
        ("bullet",
         "EVI corrects NDVI for canopy background and atmospheric "
         "effects, useful where NDVI saturates [15]."),
        ("bullet",
         "NDWI = (Green − NIR) / (Green + NIR), measuring water content "
         "and drought stress."),
        ("p",
         "GeoPandas, Rasterio, and Folium handle spatial processing. "
         "Folium is used for static maps and Leaflet (via react-leaflet) "
         "for the interactive dashboard."),

        ("h2", "3.6 Software and Tools"),
        ("p",
         "Table 3.2 summarises the complete software stack adopted in the "
         "project, organised by category. All listed tools are open-source "
         "or freely available for academic use."),
        ("table", (
            ["Category", "Tools"],
            [
                ["Programming Language", "Python 3.12, TypeScript 5"],
                ["ML Libraries", "Scikit-learn 1.5, XGBoost 3.2"],
                ["DL Frameworks", "TensorFlow 2.17 / Keras 3.14"],
                ["Geospatial", "GeoPandas, Rasterio, Folium, Shapely"],
                ["Satellite Data", "Google Earth Engine Python/JS API"],
                ["Data Processing", "Pandas, NumPy, SciPy"],
                ["Visualization", "Matplotlib, Seaborn, Plotly"],
                ["Interpretability", "SHAP (TreeExplainer, KernelExplainer)"],
                ["Web Backend", "Flask 3, flask-cors"],
                ["Frontend", "Next.js 16 (App Router), React 19, Tailwind CSS 4"],
                ["UI Components", "Radix UI, lucide-react, react-leaflet 5"],
                ["Version Control", "Git, GitHub"],
                ["IDE", "Visual Studio Code"],
            ],
            "Table 3.2: Software and tools",
        )),

        ("h2", "3.7 Evaluation Methodology"),
        ("p",
         "Four methodological technologies are adopted to ensure honest, "
         "decision-relevant performance estimates."),
        ("h3", "3.7.1 Leave-One-Year-Out Cross-Validation"),
        ("p",
         "LOYO-CV is the central evaluation protocol. For each year Y, "
         "the model is trained on every other year and tested on Y. Out-"
         "of-fold predictions are concatenated to compute metrics on "
         "data the model has never seen. This respects the temporal "
         "structure of agricultural data and prevents the leakage that "
         "would occur with a random train-test split. LOYO-CV is "
         "considered the gold standard for time-series evaluation in "
         "agricultural ML [1, 5]."),
        ("h3", "3.7.2 Statistical Significance Testing"),
        ("p",
         "Reporting a higher R² is insufficient on its own; differences "
         "must be tested for significance. Two paired tests are used: "
         "the Wilcoxon signed-rank test (non-parametric, robust to "
         "non-normal residuals) and the paired t-test (higher power when "
         "normality holds). Both are computed on paired absolute "
         "residuals, with Wilcoxon as the primary test because "
         "agricultural residuals often exhibit heavy tails."),
        ("h3", "3.7.3 SHAP Feature Attribution"),
        ("p",
         "SHAP [19] is a game-theoretic framework attributing a "
         "prediction to its input features. For each prediction, SHAP "
         "assigns a contribution value to each feature such that the "
         "sum equals the difference between the prediction and the "
         "dataset mean. Tree-based models support fast exact SHAP via "
         "TreeExplainer, used in this project. Outputs are rendered as a "
         "beeswarm summary plot and a horizontal bar chart of mean "
         "absolute SHAP values."),
        ("h3", "3.7.4 Reproducibility Tools"),
        ("p",
         "All randomness is seeded (RANDOM_STATE = 42), applied "
         "uniformly to NumPy, scikit-learn, XGBoost, and TensorFlow. The "
         "synthetic generator is deterministic. Model artefacts, plots, "
         "and result CSVs are written to outputs/ with stable filenames; "
         "the pipeline is invoked through a single main.py entry point. "
         "Git provides version control with code, data, and outputs "
         "cleanly separated."),

        ("h2", "3.8 Summary"),
        ("p",
         "This chapter has described the technologies adopted for the "
         "project, covering data sources, ML and DL algorithms, remote "
         "sensing tools, software frameworks, and evaluation "
         "methodology. These technologies are appropriate because they "
         "are freely available, well documented, and have been "
         "demonstrated effective in agricultural yield prediction "
         "research [1, 5, 8, 19]. The next chapter presents the system "
         "approach in terms of users, inputs, outputs, and the high-"
         "level processing pipeline."),
        ("pagebreak", None),
    ]


def get_chapter_4() -> list[tuple]:
    return [
        ("h1", "Chapter 4 — System Approach"),
        ("h2", "4.1 Introduction"),
        ("p",
         "The previous chapter described the technologies adopted for the "
         "project. This chapter presents how those technologies are "
         "applied to solve the big onion yield prediction problem, "
         "describing the system approach in terms of users, inputs, "
         "outputs, the five-stage processing pipeline, and the high-level "
         "system architecture."),

        ("h2", "4.2 Users"),
        ("p",
         "The system targets four user groups. (1) Department of Census "
         "and Statistics officers use model-based forecasts to cross-"
         "check field surveys and extend coverage. (2) Department of "
         "Agriculture policy planners set import quotas balancing "
         "consumer prices against farmer livelihoods. (3) District "
         "agricultural officers advise farmers using per-district "
         "predictions delivered through the choropleth map. (4) "
         "Agricultural researchers study crop-climate relationships "
         "using the SHAP feature attribution to identify which variables "
         "drive yield in the Sri Lankan context."),

        ("h2", "4.3 Inputs and Outputs"),
        ("p",
         "Inputs are organised into four feature groups derived from the "
         "Chapter 3 data sources:"),
        ("bullet",
         "Weather (9 features): rainfall, temperature, humidity, solar "
         "radiation, growing degree days, SPI drought index. Source: "
         "NASA POWER, CHIRPS."),
        ("bullet",
         "Satellite (11 features): season mean/max/min NDVI, NDVI std and "
         "anomaly, time-to-peak NDVI, NDVI growth rate, EVI, NDWI, MODIS "
         "LST day/night. Source: MODIS, Sentinel-2."),
        ("bullet",
         "Historical (5 features): previous-season yield, previous-year "
         "yield, 3-year mean, season indicator, previous-season extent. "
         "Source: DCS surveys."),
        ("bullet",
         "Soil (4 features): pH, organic carbon, clay %, sand %. "
         "District-static. Source: SoilGrids."),
        ("p",
         "Outputs are predicted yield values in MT/Ha with 95 % "
         "confidence intervals from the residual distribution. They are "
         "exposed through the Flask REST API at /predict and rendered "
         "in the dashboard as a choropleth, a result panel, and seasonal "
         "comparison charts. Auxiliary outputs include /models/compare "
         "(seven-model metrics), /feature-importance (top-15 SHAP), and "
         "/context (prediction-form prefill)."),

        ("h2", "4.4 System Pipeline"),
        ("p",
         "The system follows a five-stage pipeline, illustrated in "
         "Figure 4.2. Each stage performs a discrete responsibility and "
         "produces artefacts consumed by the next stage."),
        ("fig", (os.path.join(FIGURES_DIR, "figure_4_2_pipeline_stages.png"),
                 "Figure 4.2: System pipeline stages")),
        ("p",
         "Stage 1 — Data Acquisition. Automated collection of satellite "
         "imagery via the GEE Python API, weather data via the NASA "
         "POWER REST API, and manual extraction of yield records from "
         "DCS publications, producing seven CSV files in data/raw/."),
        ("p",
         "Stage 2 — Data Preprocessing. Data are cleaned, missing values "
         "are imputed (temporal interpolation for satellite, forward-fill "
         "for weather), cloud-contaminated pixels are masked, and all "
         "sources are aligned to Yala (April–August) and Maha (October–"
         "March) seasons. Output: data/processed/integrated_dataset.csv."),
        ("p",
         "Stage 3 — Feature Engineering. Vegetation indices, derived "
         "climate features (cumulative rainfall, GDD, SPI), and soil "
         "characteristics are integrated. Interaction features (rainfall × "
         "NDVI, temp × humidity, NDVI × LST) expose multiplicative effects "
         "to the baseline models. Output: a 32-feature matrix for ML and a "
         "multi-tensor payload for DL."),
        ("p",
         "Stage 4 — Model Training and Evaluation. Three ML and four DL "
         "models are trained under LOYO-CV with GridSearchCV "
         "hyperparameter tuning. Out-of-fold predictions are saved per "
         "model, statistical tests are run on paired residuals, an "
         "ablation study quantifies each data source, and SHAP values "
         "are computed on the best tabular model."),
        ("p",
         "Stage 5 — Visualisation and Reporting. Predictions are exposed "
         "through the Flask REST API and consumed by the Next.js "
         "dashboard, which renders a choropleth, a smart prediction form, "
         "an explainability panel, and an admin model-performance table."),

        ("h2", "4.5 High-Level System Architecture"),
        ("p",
         "Figure 4.1 shows the high-level architecture of the proposed "
         "system, comprising three modules each owned by one team member. "
         "Module 1 is the Data Pipeline, owned by Sharuja B., responsible "
         "for data acquisition, preprocessing, and feature engineering. "
         "Module 2 is the ML/DL Prediction Engine, owned by Arkam B.H.M., "
         "responsible for model training, evaluation, SHAP analysis, and "
         "API serving. Module 3 is the Visualisation and Dashboard module, "
         "owned by Shathurya P., responsible for spatial mapping, the "
         "interactive dashboard, and user-facing interfaces."),
        ("fig", (os.path.join(FIGURES_DIR, "figure_4_1_system_architecture.png"),
                 "Figure 4.1: High-level system architecture")),
        ("p",
         "The three modules are loosely coupled through well-defined "
         "interfaces. Module 1 writes processed data to data/processed/ "
         "(or PostgreSQL in production). Module 2 reads from there, "
         "writes model artefacts to outputs/models/, and serves "
         "predictions through a REST API. Module 3 consumes the API "
         "only. This separation supports independent development, "
         "simplifies testing, and isolates changes."),

        ("h2", "4.6 End-to-End Walkthrough of a Single Prediction"),
        ("p",
         "To make the pipeline concrete, this section traces a single "
         "(Matale, Yala, 2020) prediction request through the system."),
        ("p",
         "Step 1: The user selects Matale, Yala, 2020 and clicks "
         "Load context. The browser issues GET /api/backend/context, "
         "which the Next.js dev server proxies to the Flask backend."),
        ("p",
         "Step 2: The /context handler queries the in-memory pandas "
         "DataFrame loaded at startup, returns the 32 feature values as "
         "JSON, and the dashboard's FeatureInputTabs renders them in a "
         "tabbed editor for the user to review and override."),
        ("p",
         "Step 3: The user adjusts season_total_rainfall to reflect a "
         "recent forecast and clicks Predict, sending the payload to "
         "POST /api/backend/predict."),
        ("p",
         "Step 4: The /predict handler constructs the 32-element feature "
         "vector, applies the SVR scaler if needed, calls the best "
         "model's predict method, computes the 95 % CI as ±1.96 × RMSE, "
         "and returns the prediction, bounds, model name, and R²."),
        ("p",
         "Step 5: The PredictionResult component renders the response as "
         "a numeric display with a CI bar and model badge. End-to-end "
         "round trip takes ~80 ms on a development machine, dominated by "
         "the ~50 ms model.predict call for a 200-tree Random Forest."),

        ("h2", "4.7 Non-Functional Requirements"),
        ("p",
         "Beyond functional inputs and outputs, the system has several "
         "non-functional requirements:"),
        ("bullet",
         "Latency: end-to-end prediction under 500 ms on commodity "
         "hardware. Currently met (~80 ms on synthetic data)."),
        ("bullet",
         "Reproducibility: fixed random seeds and deterministic loading "
         "ensure identical results across runs."),
        ("bullet",
         "Interpretability: SHAP explainability supports non-technical "
         "users in understanding individual predictions."),
        ("bullet",
         "Maintainability: well-known frameworks (Flask, Next.js, "
         "scikit-learn, TensorFlow) and standard interfaces keep the "
         "system maintainable by a three-person team."),
        ("bullet",
         "Cost: near-zero deployment cost via free academic-tier data "
         "sources and free hosting tiers (Vercel, Fly.io / Render)."),
        ("bullet",
         "Extensibility: new districts, crops, or data sources can be "
         "added without architectural rewrites."),
        ("bullet",
         "Privacy: only aggregate district-level data; no personally "
         "identifiable farmer information."),

        ("h2", "4.8 Summary"),
        ("p",
         "This chapter has presented the system approach including the "
         "four target user groups, the multi-source inputs and per-"
         "district outputs, the five-stage processing pipeline, the "
         "three-module high-level architecture, an end-to-end "
         "walkthrough of a single prediction request, and the seven "
         "non-functional requirements that shape the system design. The "
         "next chapter provides the detailed analysis and design of "
         "each module."),
        ("pagebreak", None),
    ]


def get_chapter_5() -> list[tuple]:
    return [
        ("h1", "Chapter 5 — Analysis and Design"),
        ("h2", "5.1 Introduction"),
        ("p",
         "This chapter provides the detailed analysis and design of the "
         "proposed yield prediction system. It covers the data flow design, "
         "feature engineering strategy, the design of each ML and DL model "
         "with particular emphasis on the novel Hybrid CNN-LSTM "
         "architecture, the data pipeline design, the dashboard design, "
         "and the underlying database schema."),

        ("h2", "5.2 Data Flow Design"),
        ("p",
         "Figure 5.1 shows the data flow diagram of the system. Data flows "
         "from five external sources, through the automated ETL pipeline, "
         "into a PostgreSQL database that feeds the ML/DL prediction "
         "engine. Predictions are served via the Flask REST API to the "
         "Next.js frontend, which the four user groups consume through "
         "their browsers."),
        ("fig", (os.path.join(FIGURES_DIR, "figure_5_1_data_flow.png"),
                 "Figure 5.1: Data flow diagram")),
        ("p",
         "The data flow follows a layered design. External sources sit "
         "at the top. The ETL pipeline cleans and merges sources into a "
         "unified relational schema. The prediction engine reads merged "
         "data, trains offline, and serves online through the API. The "
         "frontend consumes the API only. This layering supports easier "
         "debugging (each layer's inputs/outputs are inspectable), "
         "independent scaling (the API can scale horizontally without "
         "changes elsewhere), and parallel development by the three "
         "module owners."),

        ("h2", "5.3 Feature Engineering Design"),
        ("p",
         "Table 5.1 summarises the features engineered from each data "
         "source. The full feature catalogue contains thirty-two columns "
         "spanning weather, satellite, historical, soil, and interaction "
         "categories. The interaction features are products of "
         "individually predictive features (such as rainfall × NDVI) "
         "that expose multiplicative effects to the tree- and kernel-"
         "based models, which would otherwise have to learn the "
         "interaction implicitly through deeper splits or higher-order "
         "kernels."),
        ("table", (
            ["Category", "Features", "Source"],
            [
                ["Weather (9)",
                 "Season avg temp, total rainfall, avg humidity, solar radiation, growing degree days, heat-stress days, drought index (SPI), diurnal temp range, max daily rainfall",
                 "NASA POWER, CHIRPS"],
                ["Satellite (11)",
                 "Season mean/max/min NDVI, NDVI std, NDVI anomaly, time-to-peak NDVI, NDVI growth rate, mean EVI, mean NDWI, LST day, LST night",
                 "MODIS MOD13Q1, Sentinel-2, MODIS LST"],
                ["Historical (5)",
                 "Previous-season yield, previous-year yield, 3-year moving mean, season indicator, previous-season extent",
                 "DCS Big Onion Surveys"],
                ["Soil (4)",
                 "Soil pH, organic carbon, clay %, sand %",
                 "SoilGrids"],
                ["Interaction (3)",
                 "Rainfall × NDVI, Temp × humidity, NDVI × LST",
                 "Derived"],
            ],
            "Table 5.1: Feature engineering summary",
        )),
        ("p",
         "Three design decisions warrant comment. First, the season "
         "indicator is a binary scalar rather than one-hot to avoid "
         "redundant degrees of freedom. Second, weather and satellite "
         "features are aggregated to seasonal summaries for the ML "
         "models (which cannot exploit sequences directly), while DL "
         "models receive both aggregated tabular features and monthly "
         "sequences. Third, the historical-yield features add a soft "
         "autoregressive component that captures district-level "
         "inertia, complementing environmental drivers."),

        ("h2", "5.4 ML Model Design"),
        ("p",
         "All three ML models are evaluated using GridSearchCV (inner "
         "loop, TimeSeriesSplit) for hyperparameter selection nested "
         "inside LOYO-CV (outer loop) for honest out-of-fold estimates."),
        ("p",
         "Random Forest hyperparameters: n_estimators in {100, 200, 500}, "
         "max_depth in {5, 10, 20, None}, min_samples_split in {2, 5, "
         "10}, min_samples_leaf in {1, 2, 4}. XGBoost: learning_rate in "
         "{0.01, 0.05, 0.1}, max_depth in {3, 5, 7}, n_estimators in "
         "{100, 300, 500}, subsample in {0.7, 0.8, 1.0}, reg_alpha in "
         "{0, 0.1, 1.0}, reg_lambda in {1.0, 1.5, 2.0}. SVR: kernel in "
         "{rbf, linear, poly}, C in {0.1, 1, 10, 100}, gamma in "
         "{scale, 0.01, 0.1}, epsilon in {0.01, 0.1, 0.5}. SVR features "
         "are standardised inside each fold to avoid leakage; tree "
         "models are scale-invariant."),

        ("h2", "5.5 DL Model Design and the Hybrid CNN-LSTM Novelty"),
        ("p",
         "Four deep learning architectures are designed. The first three "
         "(LSTM, BiLSTM, and one-dimensional CNN) serve as comparison "
         "baselines for the fourth, which is the novel Hybrid CNN-LSTM "
         "with explicit season-indicator injection."),
        ("h3", "5.5.1 LSTM Architecture"),
        ("p",
         "Two stacked LSTM layers (64 and 32 units), each followed by "
         "Dropout 0.2, then Dense(16, ReLU) and Dense(1, linear). Input "
         "shape (5, 4): five monthly timesteps × four weather variables. "
         "Trained with Adam (lr 0.001), MSE loss, EarlyStopping "
         "(patience 20) and ReduceLROnPlateau."),
        ("h3", "5.5.2 BiLSTM Architecture"),
        ("p",
         "Each LSTM layer is wrapped in Bidirectional, doubling effective "
         "units by concatenating forward and backward hidden states. "
         "Same training configuration. The BiLSTM has ~2.5 × the LSTM's "
         "parameters (77,601 vs 30,625), increasing overfitting risk on "
         "small data."),
        ("h3", "5.5.3 1D CNN Architecture"),
        ("p",
         "Input (11, 1) → Conv1D(32, k=3) → BatchNorm → MaxPool1D(2) → "
         "Conv1D(64, k=3) → GlobalAveragePooling1D → Dense(32, ReLU) → "
         "Dropout 0.2 → Dense(1, linear). The CNN treats the 11 "
         "satellite features as a 1D signal and learns local feature "
         "interactions through small filters."),
        ("h3", "5.5.4 Hybrid CNN-LSTM with Season Indicator"),
        ("p",
         "The hybrid CNN-LSTM is the central novel contribution. It has "
         "three input branches (Figure 5.2): a satellite-features CNN "
         "((11, 1) input, same Conv1D + BatchNorm + Conv1D + "
         "GlobalAveragePooling1D as the standalone CNN), a weather-"
         "sequence LSTM ((5, 4) input, stacked 64 + 32 units with "
         "dropout), and a scalar season indicator (Yala = 1, Maha = 0) "
         "passed through unchanged. The three outputs concatenate into a "
         "97-D vector (64 + 32 + 1) feeding Dense(64, ReLU) → Dropout "
         "0.3 → Dense(32, ReLU) → Dense(1, linear)."),
        ("fig", (os.path.join(FIGURES_DIR, "figure_5_2_cnn_lstm_hybrid.png"),
                 "Figure 5.2: Hybrid CNN-LSTM architecture (novelty)")),
        ("p",
         "The novelty has three elements. First, the modality-preserving "
         "structure lets the CNN learn vegetation patterns and the LSTM "
         "learn weather temporality without either modelling the other's "
         "structure. Second, the season indicator is concatenated AFTER "
         "feature extraction, allowing the branches to learn season-"
         "agnostic extractors while the final dense layers learn season-"
         "specific yield baselines — a soft multi-task formulation in "
         "one model. Third, at ~45,000 parameters with 20 % dropout the "
         "model is calibrated for the data-scarcity regime."),
        ("p",
         "Implementation uses the Keras 3 functional API in "
         "src/dl_models.py: three Input layers, two heavy branches, "
         "Concatenate, the dense stack, and tf.keras.models.Model "
         "compiled with Adam (lr 0.001) and MSE loss."),

        ("h2", "5.6 Data Pipeline Design"),
        ("p",
         "The data pipeline (owned by Sharuja B.) connects to Google "
         "Earth Engine for satellite extraction, NASA POWER for weather, "
         "and processes DCS PDFs via table extraction. Python with "
         "earthengine-api, requests, and pdfplumber. Extracted data go "
         "to a PostgreSQL database with a schema aligned by district, "
         "year, and season. Includes feature engineering, validation "
         "checks, and scheduled execution via cron in production."),

        ("h2", "5.7 Dashboard Design"),
        ("p",
         "The dashboard (owned by Shathurya P.) uses Next.js 16 (App "
         "Router), React 19, TypeScript 5, and Tailwind CSS 4. UI "
         "primitives follow the shadcn/ui pattern with hand-written "
         "wrappers on Radix UI. Four pages:"),
        ("bullet",
         "Home (/): KPI cards, interactive Leaflet choropleth coloured "
         "by predicted yield per district, and a Plotly Yala-vs-Maha "
         "comparison chart."),
        ("bullet",
         "Predict (/predict): smart form that calls /context to prefill "
         "all 32 features; users can override before submitting; result "
         "panel renders the prediction with a 95 % CI bar."),
        ("bullet",
         "Explainability (/explainability): Plotly horizontal bar of "
         "top-15 SHAP attributions plus plain-language top-5 summary."),
        ("bullet",
         "Admin (/admin): sortable model-comparison table and a Plotly "
         "RMSE/R² combo chart with target R² = 0.75 reference line."),
        ("p",
         "A sidebar global selector (district, season, year) persists "
         "via localStorage. API calls route through /api/backend/* "
         "(Next.js rewrite, no CORS), with zod schema validation at "
         "runtime."),

        ("h2", "5.8 Database Schema"),
        ("p",
         "Figure 5.3 shows the database schema. A central yield_observations "
         "fact table stores one row per (district, season, year) "
         "observation, foreign-keyed to lookup tables for districts and "
         "seasons. Three feature tables (weather_features, "
         "satellite_features, soil_features) are normalised separately to "
         "reflect their distinct update cadences: weather and satellite "
         "features are populated per observation, while soil features are "
         "district-static. A model_predictions table stores per-prediction "
         "outputs from the prediction engine, including the model name, "
         "predicted yield, confidence bounds, and timestamp."),
        ("fig", (os.path.join(FIGURES_DIR, "figure_5_3_database_schema.png"),
                 "Figure 5.3: Database schema")),
        ("p",
         "Feature tables are normalised rather than denormalised because "
         "soil features are district-static (no duplication per "
         "observation), and source-partitioned tables support "
         "incremental loading — weather and satellite ETL jobs can "
         "populate independently with the join performed at read time."),

        ("h2", "5.9 Summary"),
        ("p",
         "This chapter has provided the detailed analysis and design of "
         "all system components, including the data flow, feature "
         "engineering catalogue, ML and DL model architectures (with "
         "particular emphasis on the novel Hybrid CNN-LSTM with season "
         "indicator), the data pipeline design, the dashboard design, and "
         "the database schema. The next chapter reports on early "
         "implementation progress."),
        ("pagebreak", None),
    ]


def get_chapter_6() -> list[tuple]:
    return [
        ("h1", "Chapter 6 — Implementation"),
        ("h2", "6.1 Introduction"),
        ("p",
         "This chapter reports on the implementation progress achieved "
         "during the first half of the project period. As required by the "
         "Faculty interim-report guidelines, this chapter is descriptive "
         "rather than evaluative, with the deeper performance discussion "
         "reserved for Chapter 7. The current state of the codebase is "
         "summarised, the pipeline modules are described in turn, "
         "preliminary results obtained on a synthetic dataset are "
         "presented, and screenshots of the running system are provided."),

        ("h2", "6.2 Repository Layout and Tooling"),
        ("p",
         "The codebase is a single Git repository hosted on GitHub. "
         "Top-level directories separate the Python ML/DL pipeline "
         "(src/), the Next.js dashboard (frontend/), data files (data/"
         "raw, data/synthetic, data/processed), generated outputs "
         "(outputs/), and documentation (docs/). Python 3.12 in an "
         "isolated venv; Node 20+ with npm. All randomness is seeded "
         "(RANDOM_STATE = 42)."),
        ("ph", "Visual Studio Code window showing the repository file tree (project root)"),

        ("h2", "6.3 Data Pipeline Implementation"),
        ("p",
         "Historical yield data has been collected from DCS Big Onion "
         "Survey reports (2004–2023) for the four districts across both "
         "seasons. FAOSTAT national totals (2000–2023) provide a sanity "
         "check. A GEE JavaScript script extracts MODIS NDVI/EVI, "
         "Sentinel-2 NDVI/NDWI, CHIRPS rainfall, MODIS LST, and SMAP "
         "soil moisture for the four districts. NASA POWER weather data "
         "is retrieved for each district's coordinates."),
        ("p",
         "Sources are merged into a unified dataset aligned by district, "
         "year, and season. Missing values use temporal interpolation "
         "(satellite) and forward-fill (weather). Cloud-contaminated "
         "pixels are masked with MODIS QA and Sentinel-2 SCL. Seasonal "
         "aggregation converts monthly data to Yala/Maha averages. "
         "Feature engineering produces 25+ derived features."),
        ("p",
         "While real-data ETL is being finalised, src/data_loader.py "
         "includes a calibrated synthetic generator that reproduces the "
         "published DCS yield ranges (Matale Yala 12–24 MT/Ha, Maha "
         "6–15) and embeds realistic NDVI/rainfall/yield correlations. "
         "Deterministic given seed 42, it produces 152 observations after "
         "the historical-feature lag, validating the pipeline end-to-end "
         "before real data is finalised."),
        ("fig", (os.path.join(EDA_DIR, "yield_distribution.png"),
                 "Figure 6.1: Yield distribution by district and season (synthetic data)")),
        ("p",
         "Figure 6.1 shows the synthetic yield distribution across the "
         "four districts and two seasons. Matale shows the highest "
         "yields and Jaffna the lowest, reflecting published statistics. "
         "Yala yields are systematically higher than Maha across all "
         "districts, reflecting the bimodal monsoon pattern."),

        ("h2", "6.4 ML Model Implementation"),
        ("p",
         "The three ML models are implemented in src/ml_models.py. Each "
         "training function performs inner GridSearchCV with "
         "TimeSeriesSplit, then an outer LOYO-CV loop. Out-of-fold "
         "predictions are persisted to outputs/results/ as JSON. Each "
         "model is refit on the full dataset and serialised to "
         "outputs/models/ as joblib pickles. SVR also serialises its "
         "standard scaler for consistent inference at the API."),
        ("ph", "Visual Studio Code showing src/ml_models.py with the train_random_forest function open"),
        ("p",
         "The implementation respects LOYO-CV strictly: for each held-"
         "out year the inner GridSearchCV operates only on remaining "
         "years, the SVR scaler is fitted only on those years, and the "
         "model predicts the held-out year without ever seeing it. This "
         "matches the agronomic reality of training on past years and "
         "applying to the current one — no temporal leakage."),
        ("fig", (os.path.join(EDA_DIR, "correlation_heatmap.png"),
                 "Figure 6.2: Correlation heatmap of features with yield (synthetic data)")),

        ("h2", "6.5 DL Model Implementation"),
        ("p",
         "The four DL models are implemented in src/dl_models.py using "
         "the Keras 3 functional API. Factory functions (build_lstm, "
         "build_bilstm, build_cnn, build_cnn_lstm_hybrid) return freshly "
         "compiled models per the Chapter 5 architectures. Training "
         "uses LOYO-CV with EarlyStopping (patience 20, "
         "restore_best_weights=True), ReduceLROnPlateau (factor 0.5, "
         "patience 10), and ModelCheckpoint callbacks per fold. After "
         "LOYO, the model refits on full data and saves to outputs/"
         "models/ in .keras format."),
        ("ph", "Visual Studio Code showing src/dl_models.py with the build_cnn_lstm_hybrid function open"),
        ("p",
         "Synthetic-mode caps epochs per fold at 50 (controlled by "
         "SYNTHETIC_MODE_DL_EPOCHS in src/config.py) instead of the 200 "
         "used for real data, cutting LOYO runtime from ~1 hour to ~5 "
         "minutes on CPU during development. With real data the cap "
         "will be lifted."),

        ("h2", "6.6 Ablation Study and SHAP"),
        ("p",
         "The ablation study (src/ablation.py) runs six experiments using "
         "XGBoost as the base learner: (A) weather only, (B) satellite "
         "only, (C) historical only, (D) soil only, (E) weather + "
         "satellite, (F) all 32 features. Each subset uses the same "
         "LOYO-CV protocol; results write to outputs/results/"
         "ablation_results.csv with a bar-chart visualisation."),
        ("p",
         "SHAP analysis (src/explainer.py) uses TreeExplainer on tree-"
         "based models. Per-prediction SHAP values are computed across "
         "the dataset, mean absolute values rank features, and the top "
         "15 are persisted to outputs/results/feature_importance.json. "
         "A beeswarm summary, importance bar chart, and dependence plot "
         "for the top feature are saved as PNGs."),

        ("h2", "6.7 Flask REST API"),
        ("p",
         "The API (src/api.py) exposes six endpoints: /health, /predict, "
         "/models/compare, /feature-importance, /context, /districts. "
         "CORS is enabled via flask-cors. On startup, the API loads the "
         "best tabular model named in outputs/results/best_model_metrics."
         "json (fallback chain: XGBoost → RF → SVR) and the integrated "
         "processed dataset to support /context."),
        ("ph", "Terminal output of `PORT=5050 python src/api.py` showing the API starting and serving requests"),
        ("p",
         "/predict accepts a JSON payload with a subset of the 32 "
         "features plus district and season; missing features default "
         "to zero. It returns the predicted yield, a 95 % CI (±1.96 × "
         "RMSE), the model name, and R². /context takes district, "
         "season, year and returns the matching dataset row — or the "
         "(district, season) historical mean if the year is unseen — "
         "powering the dashboard's smart prediction form."),

        ("h2", "6.8 Dashboard Implementation"),
        ("p",
         "The dashboard (frontend/) uses Next.js 16 (App Router), React "
         "19, TypeScript 5, and Tailwind CSS 4. UI primitives in "
         "frontend/components/ui/ are hand-written on Radix UI in the "
         "shadcn/ui style. Domain components (district map, charts, "
         "prediction form) live in frontend/components/. The four pages "
         "are under frontend/app/. Cross-page state is managed through a "
         "React Context persisted to localStorage."),
        ("p",
         "The choropleth uses react-leaflet 5. Sri Lanka district "
         "polygons are extracted from geoBoundaries, filtered to the "
         "four target districts, and served as a 33 KB GeoJSON file. "
         "The map applies a viridis colour scale (4–24 MT/Ha) with "
         "hover tooltips and click handlers that update the global "
         "selection."),
        ("ph", "Dashboard home page (/) showing KPI cards, choropleth map, and Yala-vs-Maha comparison chart"),
        ("ph", "Dashboard predict page (/predict) showing the smart prediction form with all 32 features prefilled and a yield prediction"),
        ("ph", "Dashboard explainability page (/explainability) showing the top-15 SHAP feature bar chart"),
        ("ph", "Dashboard admin page (/admin) showing the sortable model comparison table and the RMSE/R² combo chart"),

        ("h2", "6.9 Preliminary Results on Synthetic Data"),
        ("p",
         "All seven models have been trained end-to-end on the synthetic "
         "dataset, with results in outputs/results/model_comparison.csv. "
         "Table 6.1 reports the preliminary metrics. These numbers are "
         "preliminary because (a) the synthetic dataset approximates "
         "DCS distributions but lacks real-world noise, and (b) DL "
         "models are trained with a 50-epoch cap rather than 200, "
         "deliberately under-trained to keep iteration fast."),
        ("table", (
            ["Model", "RMSE", "MAE", "R²", "MAPE %", "Train (s)", "Params"],
            [
                ["Random Forest", "2.067", "1.681", "0.842", "24.66", "4.78", "—"],
                ["XGBoost", "2.162", "1.756", "0.827", "25.52", "4.01", "—"],
                ["SVR", "2.868", "2.298", "0.696", "33.81", "0.07", "—"],
                ["CNN", "4.442", "3.601", "0.271", "35.49", "51.19", "8,577"],
                ["Hybrid CNN-LSTM", "4.943", "3.956", "0.097", "36.59", "83.66", "44,929"],
                ["BiLSTM", "5.101", "4.095", "0.039", "66.27", "72.58", "77,601"],
                ["LSTM", "5.148", "4.211", "0.021", "64.92", "51.74", "30,625"],
            ],
            "Table 6.1: Preliminary model comparison on synthetic data (LOYO-CV, 50 epochs/fold for DL)",
        )),
        ("fig", (os.path.join(RESULTS_PLOTS_DIR, "model_comparison_bar.png"),
                 "Figure 6.3: Model comparison: RMSE (left axis) and R² (right axis), with target R² = 0.75 dashed")),
        ("p",
         "Three observations follow. First, classical ML (RF, XGBoost) "
         "outperforms DL on this synthetic dataset — expected at 152 "
         "rows, where tree ensembles dominate. Second, Random Forest "
         "reaches R² = 0.842, exceeding the proposal's R² > 0.75 "
         "threshold. Third, the under-trained DL models are not yet "
         "competitive but their architectures are validated end-to-end "
         "and ready for full training on real data."),
        ("p",
         "The ablation yields a striking finding (Table 6.2, Figure 6.4): "
         "adding satellite features to weather improves R² by 55.3 pp "
         "(0.256 → 0.809), establishing satellite imagery as the most "
         "informative single source in the synthetic-data regime."),
        ("table", (
            ["Experiment", "Features Used", "RMSE", "MAE", "R²"],
            [
                ["A — Weather only", "9 weather features", "4.488", "3.674", "0.256"],
                ["B — Satellite only", "11 satellite features", "2.318", "1.867", "0.802"],
                ["C — Historical only", "5 historical features", "3.785", "3.041", "0.471"],
                ["D — Soil only", "4 soil features", "6.174", "5.036", "−0.408"],
                ["E — Weather + Satellite", "20 features", "2.272", "1.794", "0.809"],
                ["F — All features", "32 features", "2.154", "1.747", "0.829"],
            ],
            "Table 6.2: Ablation experiment results on synthetic data (XGBoost base learner, LOYO-CV)",
        )),
        ("fig", (os.path.join(RESULTS_PLOTS_DIR, "ablation_comparison.png"),
                 "Figure 6.4: Ablation study — feature-source contribution to R²")),
        ("p",
         "SHAP analysis identifies the interaction feature ndvi_x_lst "
         "(NDVI × LST) as the single most influential predictor, "
         "followed by season_min_ndvi, season_mean_ndvi, "
         "season_max_ndvi, and yield_3yr_avg. NDVI-based features "
         "dominating the importance ranking is consistent with the "
         "ablation's headline that satellite vegetation indices "
         "contribute the most."),
        ("fig", (os.path.join(RESULTS_PLOTS_DIR, "shap_importance.png"),
                 "Figure 6.5: SHAP top-15 feature importance (best tabular model)")),

        ("h2", "6.10 Summary"),
        ("p",
         "Interim implementation covers the complete ML/DL pipeline "
         "(seven models, LOYO-CV, ablation, SHAP, statistical tests), a "
         "Flask REST API with six endpoints, and a four-page Next.js "
         "dashboard. The pipeline has been validated end-to-end on "
         "synthetic data, achieving R² = 0.842 with Random Forest. "
         "Real-data ingestion is the primary remaining task. The next "
         "chapter discusses how this work differs from existing "
         "approaches, challenges encountered, and the plan for the "
         "second half of the project."),
        ("pagebreak", None),
    ]


def get_chapter_7() -> list[tuple]:
    return [
        ("h1", "Chapter 7 — Discussion and Further Work"),
        ("h2", "7.1 Introduction"),
        ("p",
         "This chapter discusses the current status of the project, the "
         "ways in which the proposed solution differs from existing work, "
         "the challenges encountered during the first half of the project "
         "period, and the plan for the remaining work."),

        ("h2", "7.2 How This Work Differs from Others"),
        ("p",
         "This project differs from existing work in five complementary "
         "ways."),
        ("p",
         "First, it is the first AI-powered yield prediction system for "
         "big onion in South Asia. Existing onion ML work has focused "
         "on different geographies (Bangladesh [9], Korea [11]) or "
         "different outcomes (price rather than yield in Sri Lanka [14])."),
        ("p",
         "Second, it integrates four data streams (satellite, weather, "
         "soil, historical yield) that have never been combined for a "
         "vegetable crop. Existing work uses weather alone [5, 9, 10] or "
         "weather plus a single satellite modality [7]. The four-stream "
         "fusion enables the ablation decomposition in Chapter 6."),
        ("p",
         "Third, it provides a systematic ML-versus-DL comparison for "
         "vegetable yield with scarce data. Cereal-crop comparisons exist "
         "[3, 7] but the small-data bimodal-monsoon regime is "
         "qualitatively different. The project answers whether DL is "
         "justified at this data scale, using LOYO-CV and Wilcoxon / "
         "paired t-tests for rigour."),
        ("p",
         "Fourth, the Hybrid CNN-LSTM introduces explicit season-"
         "indicator injection not present in cited hybrid yield models "
         "[8]. The indicator is concatenated AFTER the CNN/LSTM "
         "branches' feature extraction, letting the model learn season-"
         "agnostic extractors while final dense layers learn season-"
         "specific baselines. This soft multi-task formulation is "
         "engineered for the Yala/Maha bimodal system and is novel for "
         "crop yield prediction."),
        ("p",
         "Fifth, district-level spatial predictions on an interactive "
         "choropleth address the gap that current Sri Lankan studies "
         "produce only national aggregates. This transforms the model "
         "from research artefact into usable decision-support."),

        ("h2", "7.3 Challenges Encountered"),
        ("p",
         "Three categories of challenge have been encountered."),
        ("p",
         "The methodological challenge is data scarcity. With ~160 "
         "observations (4 districts × 2 seasons × 20 years), the dataset "
         "is two to three orders of magnitude smaller than typical "
         "cereal-crop DL datasets. This shapes every modelling decision: "
         "small deep models (8K–80K parameters), aggressive "
         "regularisation, and LOYO-CV instead of random splits. The "
         "Chapter 6 results confirm classical ML outperforms DL at this "
         "scale, validating the seven-model comparison."),
        ("p",
         "The data-engineering challenge is that DCS reports are PDFs, "
         "not structured datasets. A semi-automated pipeline using "
         "pdfplumber has been developed but each extracted table still "
         "requires manual verification. A second challenge is cloud "
         "contamination in optical satellite imagery; MODIS QA bands "
         "and Sentinel-2 SCL mask contaminated pixels but this reduces "
         "effective temporal resolution."),
        ("p",
         "The engineering challenge is coordinating three independently "
         "developed modules in different languages and processes. "
         "Communicating exclusively through well-defined interfaces "
         "(filesystem CSVs/JSON for pipeline↔engine, REST API for "
         "engine↔dashboard) reduces interaction cost and enables "
         "parallel development."),

        ("h2", "7.4 Further Work Plan"),
        ("p",
         "The remaining work is organised into eight workstreams:"),
        ("bullet",
         "Real-data validation. Re-run the pipeline on real CSVs, lift "
         "the synthetic-mode epoch cap, and run the full hyperparameter "
         "grids. Compare against synthetic-data baselines."),
        ("bullet",
         "Statistical significance reporting. Publish Wilcoxon and "
         "paired t-test p-values for every pairwise comparison and "
         "aggregate ML-vs-DL with statistical rigour."),
        ("bullet",
         "Stacking ensemble. A meta-learner combining best ML and best "
         "DL predictions through regularised meta-regression, on the "
         "expectation that ML and DL errors are partially uncorrelated."),
        ("bullet",
         "Spatial error analysis. Decompose errors by district, "
         "identifying biases, and explore district embeddings. The "
         "synthetic data shows Matale easiest, Anuradhapura hardest."),
        ("bullet",
         "Early prediction feasibility. Train LSTM and Hybrid on "
         "truncated weather sequences (end of months 1, 2, 3) and "
         "evaluate prediction quality vs lead time."),
        ("bullet",
         "Cloud deployment. Deploy the Flask API to Fly.io or Render "
         "and the Next.js frontend to Vercel with DNS, SSL, and basic "
         "monitoring for stakeholder access during the demonstration."),
        ("bullet",
         "User testing. Five to ten usability sessions with the four "
         "target user groups, collecting feedback on the choropleth, "
         "prediction form, and explainability panel."),
        ("bullet",
         "Final documentation and demonstration. Finalise the report "
         "and prepare a live demonstration with recorded backup."),

        ("h2", "7.5 Methodological Choices and Their Justifications"),
        ("p",
         "Four methodological choices warrant explicit justification."),
        ("h3", "7.5.1 Why Leave-One-Year-Out Rather Than k-Fold"),
        ("p",
         "Random k-fold is inappropriate for time-series — a model "
         "fitted on a fold containing both 2010 and 2020 observations "
         "has effectively seen the future when predicting intermediate "
         "years. LOYO-CV avoids this leakage. The cost is computational: "
         "twenty trainings instead of five. For tree models this is "
         "negligible; for DL it is significant, motivating the synthetic-"
         "mode epoch cap in Chapter 6."),
        ("h3", "7.5.2 Why Synthetic Data Calibrated to Published Distributions"),
        ("p",
         "The benefit is that the pipeline is implemented, tested, and "
         "benchmarked in advance — when real data arrives the only "
         "operation required is to drop the CSVs and re-run main.py. "
         "The cost is that synthetic results are not proper estimates "
         "for the real-data system: the generator embeds clean NDVI/"
         "yield and rainfall/yield correlations that are simpler than "
         "real data. We expect real-data R² to be lower, possibly "
         "substantially so for the deep models."),
        ("h3", "7.5.3 Why a Hybrid CNN-LSTM and Not a Transformer"),
        ("p",
         "Transformers dominate general sequence modelling but require "
         "an order of magnitude more parameters than LSTMs and more data "
         "to avoid overfitting. With 152 training observations the "
         "hybrid CNN-LSTM is already at the upper bound of what can be "
         "responsibly trained; a transformer would almost certainly "
         "underperform. As data scales (more districts, finer spatial "
         "granularity), transformers become a reasonable extension."),
        ("h3", "7.5.4 Why ML and DL Together Rather Than One or the Other"),
        ("p",
         "A purely ML project would not address whether DL is justified "
         "for vegetable yield with scarce data. A purely DL project "
         "would lack baselines to make claims credible. Including both "
         "under the same protocol with statistical significance testing "
         "is what transforms the comparison from heuristic preference "
         "into a defensible empirical finding."),

        ("h2", "7.6 Limitations and Risks"),
        ("p",
         "The most significant limitation is dataset size: no "
         "methodological improvement can fully compensate for the small "
         "number of observations, so real-data DL performance is "
         "uncertain. Mitigation: the multi-model approach guarantees a "
         "useful system even if DL underperforms — the strongest "
         "classical model remains the production candidate."),
        ("p",
         "A second limitation is satellite data quality. Cloud cover "
         "during monsoon transitions can produce extended NDVI gaps. "
         "Mitigation: combining MODIS and Sentinel-2 (different revisit "
         "cadences) reduces simultaneous contamination, and the SPI "
         "drought index provides a precipitation-based alternative."),
        ("p",
         "A third risk is institutional acceptance by DCS and the "
         "Department of Agriculture, which produce statistics through "
         "long-established manual procedures. Mitigation: the system is "
         "positioned as decision-support rather than replacement, SHAP "
         "supports human review of every prediction, and the "
         "supervisor's institutional network provides a channel for "
         "early stakeholder engagement."),

        ("h2", "7.7 Summary"),
        ("p",
         "The project has made significant progress in the first half of "
         "the project period, with a complete ML/DL modelling pipeline, "
         "a Flask REST API, and a Next.js dashboard all implemented and "
         "validated on a synthetic dataset. Five distinct novel "
         "contributions have been identified, three categories of "
         "challenge have been documented, four methodological choices "
         "have been explicitly justified, three known limitations have "
         "been articulated with mitigations, and an eight-workstream "
         "plan for the remaining period has been laid out. The next "
         "section presents the alphabetised list of references cited "
         "throughout this report."),
        ("pagebreak", None),
    ]


# ---------- References (alphabetical, Karunananda style) --------------------

REFERENCES = [
    # Each entry is one paragraph, formatted as the guideline Appendix C example.
    "[1] Jabed, M. A., Karim, M. R., Hossain, M. S., Hossain, M. K., Khan, M. M. and Begum, M. (2024), Crop yield prediction in agriculture: A comprehensive review of ML and DL approaches, with insights for future research, Heliyon, 10(24), e40836.",
    "[2] Kamilaris, A. and Prenafeta-Boldú, F. X. (2018), Deep learning in agriculture: A survey, Computers and Electronics in Agriculture, 147, pp. 70–90.",
    "[3] Chikwendu, F., Ahmed, A. and Ortega-Mansilla, A. (2025), A comparative study of machine learning models in predicting crop yield, Discover Agriculture, Springer.",
    "[4] Department of Census and Statistics Sri Lanka (2010–2025), Big Onion Survey Reports. Available: https://www.statistics.gov.lk",
    "[5] Amarasinghe, A., De Silva, P. and Wijesinghe, K. (2024), Advancing food sustainability: rice yield prediction in Sri Lanka using weather-based, feature-engineered ML models, Discover Applied Sciences, 6, 603.",
    "[6] Department of Agriculture Sri Lanka (n.d.), HORDI Crop – Big Onion: Cultivation Guidelines. Available: https://doa.gov.lk",
    "[7] Wang, Z., Yan, Y., Lu, Q. and Yang, P. (2022), Machine learning crop yield models based on meteorological features and comparison with a process-based model, Artificial Intelligence for the Earth Systems, 1(4), pp. 1–18.",
    "[8] Rajpoot, S. and Chandrakar, O. (2025), A hybrid CNN-LSTM deep learning framework for enhanced crop yield prediction using spatial-temporal agricultural data, International Journal of Statistics and Applied Mathematics, SP-10(12), pp. 1–9.",
    "[9] Iqbal, D. M., Hossain, S., Rahman, M. S. and Karim, R. (2023), OnionBangla: A supervised ML approach for predicting onion yield using Bangladeshi climate data, in Proceedings of IEEE ICCCI, pp. 178–185.",
    "[10] Wickramasinghe, L., Weliwita, J. A. and Ekanayake, I. U. (2021), Modeling the relationship between rice yield and climate variables using statistical and ML techniques, Journal of Mathematics, 2021, 1972927.",
    "[11] Kim, W. and Soon, B. M. (2023), Advancing agricultural predictions: A deep learning approach to estimating bulb weight using Neural Prophet model, Agronomy, 13(5), 1305.",
    "[12] Herath, K., Perera, B. and Hewage, S. (2016), Extraction of agricultural phenological parameters of Sri Lanka using MODIS NDVI time series data, Procedia Food Science, 6, pp. 1248–1254.",
    "[13] FAOSTAT (2024), Food and Agriculture Statistics — Crops and Livestock Products. Available: https://www.fao.org/faostat",
    "[14] Sutharsan, J. and Yogendran, S. (2023), Red onion price factors correlation identification and price prediction using multiple ML models for Jaffna district Sri Lanka, in Proceedings of the IEEE Conference on Computational Intelligence, pp. 412–419.",
    "[15] MODIS Science Team (2019), MODIS Vegetation Index User Guide, NASA. Available: https://modis.gsfc.nasa.gov",
    "[16] Hochreiter, S. and Schmidhuber, J. (1997), Long short-term memory, Neural Computation, 9(8), pp. 1735–1780.",
    "[17] Breiman, L. (2001), Random forests, Machine Learning, 45(1), pp. 5–32.",
    "[18] Chen, T. and Guestrin, C. (2016), XGBoost: A scalable tree boosting system, in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785–794.",
    "[19] Lundberg, S. M. and Lee, S.-I. (2017), A unified approach to interpreting model predictions, in Advances in Neural Information Processing Systems 30, pp. 4765–4774.",
    "[20] LeCun, Y., Bengio, Y. and Hinton, G. (2015), Deep learning, Nature, 521(7553), pp. 436–444.",
    "[21] Pedregosa, F., Varoquaux, G., Gramfort, A. and others (2011), Scikit-learn: Machine learning in Python, Journal of Machine Learning Research, 12, pp. 2825–2830.",
    "[22] Abadi, M., Barham, P., Chen, J., Chen, Z., Davis, A. and others (2016), TensorFlow: A system for large-scale machine learning, in Proceedings of the 12th USENIX Symposium on Operating Systems Design and Implementation, pp. 265–283.",
    "[23] Gorelick, N., Hancher, M., Dixon, M., Ilyushchenko, S., Thau, D. and Moore, R. (2017), Google Earth Engine: Planetary-scale geospatial analysis for everyone, Remote Sensing of Environment, 202, pp. 18–27.",
    "[24] Hengl, T., Mendes de Jesus, J., Heuvelink, G. B. M. and others (2017), SoilGrids250m: Global gridded soil information based on machine learning, PLoS ONE, 12(2), e0169748.",
    "[25] Funk, C., Peterson, P., Landsfeld, M., Pedreros, D., Verdin, J. and others (2015), The climate hazards infrared precipitation with stations — a new environmental record for monitoring extremes, Scientific Data, 2, 150066.",
    "[26] Stackpole, B. (n.d.), NASA POWER project — Prediction of Worldwide Energy Resources. Available: https://power.larc.nasa.gov",
    "[27] Drusch, M., Del Bello, U., Carlier, S. and others (2012), Sentinel-2: ESA's optical high-resolution mission for GMES operational services, Remote Sensing of Environment, 120, pp. 25–36.",
    "[28] Wan, Z., Hook, S. and Hulley, G. (2015), MOD11A1 MODIS/Terra land surface temperature/emissivity daily L3 global 1km SIN grid V006, NASA EOSDIS Land Processes DAAC.",
    "[29] Karunananda, A. S. (2006), Guidelines for preparation of Interim Reports, Faculty of Information Technology, University of Moratuwa, internal document.",
    "[30] Pearson, K. (1895), Notes on regression and inheritance in the case of two parents, Proceedings of the Royal Society of London, 58, pp. 240–242.",
]


def add_references(doc: Document) -> None:
    _heading(doc, "References", level="chapter")
    _para(
        doc,
        "Citations in the body of this report follow the Karunananda interim "
        "report guideline format using square-bracketed numbers. The list "
        "below is sorted alphabetically by the surname of the first author "
        "of each work and renumbered. Each citation in the text refers to "
        "the corresponding entry by its bracketed number.",
    )
    # Sort REFERENCES alphabetically by the first author surname and renumber.
    def _sort_key(entry: str) -> str:
        # The entry starts "[N] Surname, F.M., ..." — take Surname.
        body = entry.split("] ", 1)[1] if "] " in entry else entry
        return body.split(",")[0].strip().lower()

    sorted_refs = sorted(REFERENCES, key=_sort_key)
    for i, ref in enumerate(sorted_refs, start=1):
        # Replace the original [n] prefix with the new index.
        text = ref.split("] ", 1)[1] if "] " in ref else ref
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        run = p.add_run(f"[{i}] {text}")
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
    _page_break(doc)


# ---------- Appendices ------------------------------------------------------

def add_appendix_a(doc: Document) -> None:
    _heading(doc, "Appendix A — Individuals' Contribution to the Project", level="chapter")
    _para(
        doc,
        "This appendix records the individual contribution of each group "
        "member to the project during the first half of the project period, "
        "as required by the Faculty interim-report guidelines.",
    )

    # ---- Arkam ----
    _heading(doc, "A.1 Arkam B.H.M. (214019K)", level="section")
    _para(
        doc,
        "My contribution focuses on the Machine Learning, Deep Learning, "
        "and Modelling Research component. During the first half I "
        "conducted a literature review of 30+ papers on ML/DL crop yield "
        "prediction, identifying the five research gaps in Chapter 2. I "
        "designed and implemented the complete modelling pipeline: three "
        "ML models (RF, XGBoost, SVR) and four DL models (LSTM, BiLSTM, "
        "1D CNN, and the novel Hybrid CNN-LSTM with season-indicator "
        "injection). All seven are evaluated using LOYO-CV with inner-"
        "loop GridSearchCV (TimeSeriesSplit) for hyperparameter tuning. "
        "The novel Hybrid CNN-LSTM is the central research contribution: "
        "parallel CNN (satellite) and LSTM (weather) branches with the "
        "season indicator concatenated after feature extraction, so the "
        "final dense layers learn season-specific baselines while "
        "sharing earlier feature extraction across seasons.",
    )
    _para(
        doc,
        "Beyond the core models I implemented a six-experiment ablation "
        "study, SHAP feature attribution (TreeExplainer), statistical "
        "tests (Wilcoxon, paired t), and a Flask REST API with six "
        "endpoints. The pipeline is validated end-to-end on synthetic "
        "data calibrated to published DCS yield ranges, with Random "
        "Forest reaching R² = 0.842 — exceeding the R² > 0.75 threshold. "
        "The principal challenge was preventing overfitting on the small "
        "dataset (~160 rows); addressed through shallow architectures, "
        "dropout, early stopping, and rigorous LOYO-CV. I learned that "
        "for small agricultural datasets, feature engineering and "
        "evaluation methodology often matter more than model complexity. "
        "Remaining work: real-data validation, stacking ensemble, and "
        "the final report's methodology, results, and discussion "
        "chapters.",
    )

    # ---- Sharuja ----
    _heading(doc, "A.2 Sharuja B. (214192G)", level="section")
    _para(
        doc,
        "My contribution focuses on Data Engineering, Feature Engineering, "
        "and Data-Centric Research. During the first half I collected "
        "historical DCS Big Onion Survey data (2004–2023) by extracting "
        "tables from PDFs using pdfplumber, with hand-verification. I "
        "developed a Google Earth Engine JavaScript script that automates "
        "extraction of MODIS NDVI/EVI 16-day composites, Sentinel-2 "
        "NDVI/NDWI monthly composites, CHIRPS rainfall, MODIS LST "
        "(day/night), and SMAP soil moisture for the four target "
        "districts. I implemented NASA POWER weather retrieval (Python "
        "requests library) for each district's coordinates, and collected "
        "SoilGrids soil characteristics (pH, organic carbon, clay %, "
        "sand %) via REST API.",
    )
    _para(
        doc,
        "I designed the preprocessing workflow: temporal alignment to "
        "Yala and Maha seasons, cloud masking via MODIS QA / Sentinel-2 "
        "SCL, missing-value imputation (temporal interpolation for "
        "satellite, forward-fill for weather), and normalisation. I "
        "engineered 25+ derived features including growing degree days, "
        "cumulative rainfall, SPI, NDVI anomaly, time-to-peak NDVI, "
        "NDVI growth rate, and interaction features. The principal "
        "challenge was aligning sources at different temporal "
        "resolutions (daily weather, 16-day satellite, seasonal yield); "
        "I aggregated to seasonal for ML and retained monthly for the "
        "LSTM input. I learned that data engineering is often the "
        "majority of an AI project and that careful schema design pays "
        "off when multiple consumers read the same dataset. Remaining "
        "work: NDVI vs EVI vs NDWI comparison, growth-stage analysis, "
        "MODIS vs Sentinel-2 resolution comparison, and the production "
        "database automation pipeline.",
    )

    # ---- Shathurya ----
    _heading(doc, "A.3 Shathurya P. (214193K)", level="section")
    _para(
        doc,
        "My contribution focuses on System Integration, Visualisation, "
        "and Decision-Support Research. During the first half I designed "
        "the overall system architecture and the loosely coupled file-"
        "system and REST API interfaces between the three modules. I "
        "created wireframes for the four dashboard pages and reviewed "
        "them with stakeholders before implementation. I implemented the "
        "full Next.js 16 dashboard (App Router, React 19, TypeScript 5, "
        "Tailwind CSS 4), including hand-written UI primitives in the "
        "shadcn/ui style on Radix UI.",
    )
    _para(
        doc,
        "I implemented the four pages: home (KPI cards, Leaflet "
        "choropleth, Plotly seasonal comparison); predict (smart form "
        "calling /context to prefill all 32 features, with prediction "
        "and CI rendering); explainability (top-15 SHAP horizontal bar "
        "with plain-language top-5 explanations); and admin (sortable "
        "model table, Plotly RMSE/R² combo chart with R² > 0.75 "
        "reference). I sourced and filtered the Sri Lanka district "
        "GeoJSON from geoBoundaries (25 districts/241 KB → 4 districts/"
        "33 KB) and researched ensemble fusion strategies for the next "
        "phase. The principal challenge was designing a dashboard "
        "accessible to non-technical agricultural stakeholders; I "
        "addressed this by separating the prediction view (single "
        "number + CI) from the explanation view (SHAP chart), and "
        "labelling every input with a plain-language description and "
        "units. I learned that effective decision-support requires as "
        "much attention to information hierarchy and progressive "
        "disclosure as to model accuracy. Remaining work: spatial "
        "error analysis, early-prediction feasibility, cloud "
        "deployment, and user testing with five to ten participants.",
    )

    # Distribution table
    _heading(doc, "A.4 Distribution of Work", level="section")
    _para(
        doc,
        "Table A.1 summarises the headline distribution of work across "
        "the three group members at the end of the first half of the "
        "project period.",
    )
    _add_table(
        doc,
        ["Workstream", "Owner", "Status"],
        [
            ["Literature review", "Arkam B.H.M.", "Complete"],
            ["Data acquisition (DCS, NASA POWER, GEE, SoilGrids)", "Sharuja B.", "Complete (real-data ETL ongoing)"],
            ["Data preprocessing and feature engineering", "Sharuja B.", "Complete on synthetic; real-data integration ongoing"],
            ["ML model implementation (RF, XGBoost, SVR)", "Arkam B.H.M.", "Complete (LOYO-CV, GridSearchCV)"],
            ["DL model implementation (LSTM, BiLSTM, CNN, Hybrid)", "Arkam B.H.M.", "Complete (synthetic-mode)"],
            ["Ablation study (six experiments)", "Arkam B.H.M.", "Complete on synthetic"],
            ["SHAP feature attribution", "Arkam B.H.M.", "Complete on synthetic"],
            ["Statistical significance testing", "Arkam B.H.M.", "Complete on synthetic"],
            ["Flask REST API (six endpoints)", "Arkam B.H.M.", "Complete and verified"],
            ["System architecture design", "Shathurya P.", "Complete"],
            ["Next.js dashboard (four pages)", "Shathurya P.", "Complete with screenshot placeholders"],
            ["Choropleth GeoJSON sourcing and filtering", "Shathurya P.", "Complete"],
            ["Real-data validation", "All members", "Pending"],
            ["Stacking ensemble", "Arkam B.H.M.", "Pending"],
            ["Cloud deployment", "Shathurya P.", "Pending"],
            ["User testing", "Shathurya P.", "Pending"],
            ["Final report writing", "All members", "In progress"],
        ],
        "Table A.1: Distribution of individual contributions and current status",
    )
    _page_break(doc)


def add_appendix_b(doc: Document) -> None:
    _heading(doc, "Appendix B — Code and Implementation Excerpts", level="chapter")
    _para(
        doc,
        "This appendix collects screenshots and code excerpts referenced "
        "in Chapter 6. Each item is labelled with its source location in "
        "the project repository.",
    )

    _heading(doc, "B.1 Repository Structure", level="section")
    _para(doc, "The top-level directory layout of the project repository:")
    _placeholder(doc, "Visual Studio Code Explorer pane showing the repository top-level directories (data/, docs/, frontend/, outputs/, src/, requirements.txt, main.py, README.md)")

    _heading(doc, "B.2 Hybrid CNN-LSTM Implementation (src/dl_models.py)", level="section")
    _para(doc, "The Keras 3 functional-API definition of the novel Hybrid CNN-LSTM architecture:")
    _placeholder(doc, "Visual Studio Code window showing src/dl_models.py with the build_cnn_lstm_hybrid function visible (approximately lines 73–97)")

    _heading(doc, "B.3 Configuration File (src/config.py)", level="section")
    _para(doc, "Project-wide hyperparameters live in a single config file so all experiment knobs are visible in one place:")
    _placeholder(doc, "Visual Studio Code window showing src/config.py with the ALL_FEATURES list and the DL hyperparameters block visible")

    _heading(doc, "B.4 Flask API (src/api.py)", level="section")
    _para(doc, "The serving layer exposes six endpoints; the /predict endpoint and the new /context endpoint are particularly relevant to the dashboard's smart prediction form:")
    _placeholder(doc, "Visual Studio Code window showing src/api.py with the /predict and /context route handlers visible")

    _heading(doc, "B.5 Dashboard Home Page (frontend/app/page.tsx)", level="section")
    _para(doc, "The Next.js home page composes the KPI cards, the choropleth map, and the seasonal comparison chart. It calls the /predict endpoint for each (district, season) combination of the selected year:")
    _placeholder(doc, "Visual Studio Code window showing frontend/app/page.tsx with the predictAllDistrictsForYear helper and the JSX layout visible")

    _heading(doc, "B.6 Pipeline Smoke-Test Output", level="section")
    _para(doc, "Running `python main.py --skip-shap --skip-dl` produces the following terminal output, demonstrating the end-to-end pipeline executes successfully on the synthetic dataset and produces all expected artefacts:")
    _placeholder(doc, "Terminal window showing the output of `python main.py --skip-shap --skip-dl`, ending with 'PIPELINE COMPLETE'")


# ---------- Build the document ---------------------------------------------

def build() -> str:
    doc = Document()
    _configure_document(doc)

    # SECTION 1: cover page (no page number)
    add_cover_page(doc)

    # SECTION 2: title page (no page number) — keep within section 1, suppress via different first page
    # Actually, simpler approach: title page is page 2 of section 1, and the
    # whole section 1 has no footer. Pre-pages start in section 2.
    add_title_page(doc)

    # Section break -> Roman numerals for pre-pages
    _section_break(doc)
    pre_section = doc.sections[-1]
    # Different first page = False so all pages in this section get the page number
    pre_section.different_first_page_header_footer = False
    _add_centered_footer_pagenum(pre_section, fmt="lowerRoman", start=1)
    add_pre_pages(doc)

    # Section break -> Arabic numerals for body
    _section_break(doc)
    body_section = doc.sections[-1]
    body_section.different_first_page_header_footer = False
    _add_centered_footer_pagenum(body_section, fmt="decimal", start=1)

    # Body chapters
    _emit(doc, get_chapter_1())
    _emit(doc, get_chapter_2())
    _emit(doc, get_chapter_3())
    _emit(doc, get_chapter_4())
    _emit(doc, get_chapter_5())
    _emit(doc, get_chapter_6())
    _emit(doc, get_chapter_7())

    # References + appendices (continue Arabic numbering)
    add_references(doc)
    add_appendix_a(doc)
    add_appendix_b(doc)

    # Section 1 (cover + title) gets no footer page number — set explicitly.
    cover_section = doc.sections[0]
    _set_section_page_numbering(cover_section, fmt="decimal", start=1)
    # Hide page number on cover/title pages by leaving the footer empty.
    # different_first_page_header_footer=True suppresses on the first page of the section.
    # We additionally clear all footers explicitly:
    cover_section.first_page_footer.is_linked_to_previous = False
    if cover_section.first_page_footer.paragraphs:
        cover_section.first_page_footer.paragraphs[0].text = ""
    cover_section.footer.is_linked_to_previous = False
    if cover_section.footer.paragraphs:
        cover_section.footer.paragraphs[0].text = ""

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build()
    print(f"Wrote interim report → {out}")
    print(f"Size: {os.path.getsize(out):,} bytes")
